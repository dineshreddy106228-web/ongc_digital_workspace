"""CSC Drafting Workspace — Word document export.

Produces TWO documents in one .docx:

  PART A — ONGC Corporate Specification Sheet
            Matches CSC Format A exactly:
            • ONGC logo + "CORPORATE SPECIFICATIONS FOR OIL FIELD CHEMICALS" header
            • Double-line page border (DrawingML: thin outer 0.75pt + thick inner 3pt)
            • 4-row header info table (Chemical Name / Spec No / Test Procedure / Material Code)
              — 2 cols, 0.5pt single borders, col widths 2.58" / 4.02"
            • Parameter table: 4 cols — S.No. | Parameter Name (Type) | : | Required Value
              — TableGrid style, Times New Roman 11pt, no shading
            • Parameter type shown in brackets after the parameter name

  PART B — CSC Draft Note
            Committee sections: Background, Issues, Impact Analysis,
            Recommendation, with a formal header.

Entry point:
    build_word_document(draft_dict, sections_dict, parameters_list, flags_list=None, impact_dict=None, logo_path=None) -> bytes
    build_flask_review_document(review_context, logo_path=None) -> bytes
    build_master_spec_document(all_specs, include_draft_note=False, logo_path=None) -> bytes
"""

from __future__ import annotations

import base64
import io
import logging
import re
from datetime import datetime
from itertools import count
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor, Emu, Cm, Twips

from app.core.services.csc_utils import (
    SECTION_BACKGROUND,
    SECTION_EXISTING_SPEC,
    SECTION_PROPOSED,
    SECTION_JUSTIFICATION,
    REC_MAIN_KEY,
    deserialize_impact_checklist_state,
    summarize_impact_checklist_state,
    REC_REMARKS_KEY,
    ISSUE_TYPES,
    SPEC_SUBSET_LABELS,
    SPEC_SUBSET_ORDER,
    parse_spec_number,
    spec_sort_key,
    get_impact_grade,
)

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Colours — match the PDF exactly
# ---------------------------------------------------------------------------
_BLACK      = RGBColor(0x00, 0x00, 0x00)
_DARK_GRAY  = RGBColor(0x1F, 0x29, 0x37)
_COPPER     = RGBColor(0xA5, 0x4B, 0x23)
_COPPER_HEX = "A54B23"
_LGRAY_HEX  = "E5E7EB"
_WHITE_HEX  = "FFFFFF"

# Fonts matching the original PDF
_FONT_MAIN  = "Times New Roman"
_FONT_TITLE = "Times New Roman"

# ---------------------------------------------------------------------------
# Page border geometry — drawn from CSC Format A measurements
#
# The Format A docx uses a DrawingML wpg:wgp (WordprocessingGroup) anchored
# behind the text with behindDoc="1".  It produces a double-line frame:
#   • Outer thin line  : 10,668 EMU ≈ 0.75pt
#   • Inner thick line : 38,100 EMU ≈ 3pt
# Gap between them    : ~30,000 EMU ≈ 2pt
#
# We replicate this as an anchored DrawingML group inserted into an explicit
# paragraph anchor per spec page. All coordinates are in EMU.
# ---------------------------------------------------------------------------
_BORDER_THIN_W  = 10668    # 0.75pt — outer thin line
_BORDER_THICK_W = 38100    # 3pt    — inner thick line
_BORDER_GAP     = 28575    # 2pt gap between thin and thick

# Page content area used for the border group (A4, 0.9" margins each side)
# Width:  (8.27in - 0.9in - 0.59in) * 914400 ≈ 6,288,000 EMU
# Height: (11.69in - 0.85in - 0.19in) * 914400 ≈ 9,720,000 EMU
_BORDER_CX = 6_288_000
_BORDER_CY = 9_033_468
_BORDER_POS_X = 763_675
_BORDER_POS_Y = 914_401

_DRAWING_ID_COUNTER = count(9001)


def _next_drawing_id() -> int:
    """Generate deterministic unique drawing ids in this process."""
    return next(_DRAWING_ID_COUNTER)


def _make_border_drawing_xml(drawing_id: int) -> str:
    """Return the XML for a DrawingML double-line page border group.

    The group is positioned behind the document text (behindDoc=1) and
    contains 8 line shapes: 4 thin outer lines + 4 thick inner lines.
    All lines use round caps (a:round) matching the Format A source.
    """
    cx = _BORDER_CX
    cy = _BORDER_CY
    thin  = _BORDER_THIN_W
    thick = _BORDER_THICK_W
    gap   = _BORDER_GAP

    def ln(w: int) -> str:
        return (
            f'<a:ln w="{w}" cap="rnd" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
            f'<a:solidFill><a:srgbClr val="000000"/></a:solidFill>'
            f'</a:ln>'
        )

    def shape(sid: int, x: int, y: int, scx: int, scy: int, line_w: int) -> str:
        return f"""
        <wps:wsp xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
          <wps:cNvSpPr><a:spLocks noChangeArrowheads="1"/></wps:cNvSpPr>
          <wps:spPr>
            <a:xfrm><a:off x="{x}" y="{y}"/><a:ext cx="{scx}" cy="{scy}"/></a:xfrm>
            <a:prstGeom prst="line"><a:avLst/></a:prstGeom>
            <a:noFill/>
            {ln(line_w)}
          </wps:spPr>
          <wps:bodyPr/>
        </wps:wsp>"""

    # 8 shapes: top-thin, top-thick, bottom-thin, bottom-thick,
    #           left-thin, left-thick, right-thin, right-thick
    # Horizontal lines: cy=0 (line height), cx=full width
    # Vertical lines  : cx=0 (line width), cy=full height
    shapes = (
        # Top outer thin
        shape(1, 0,           0,           cx, 0,  thin)
        # Top inner thick
        + shape(2, gap,        gap,          cx - gap*2, 0,  thick)
        # Bottom outer thin
        + shape(3, 0,          cy,          cx, 0,  thin)
        # Bottom inner thick
        + shape(4, gap,        cy - gap,    cx - gap*2, 0,  thick)
        # Left outer thin
        + shape(5, 0,           0,           0,  cy, thin)
        # Left inner thick
        + shape(6, gap,         gap,          0,  cy - gap*2, thick)
        # Right outer thin
        + shape(7, cx,          0,           0,  cy, thin)
        # Right inner thick
        + shape(8, cx - gap,    gap,          0,  cy - gap*2, thick)
    )

    xml = f"""<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <wp:anchor distT="0" distB="0" distL="0" distR="0"
             simplePos="0" relativeHeight="251658240" behindDoc="1"
             locked="0" layoutInCell="1" allowOverlap="0"
             xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">
    <wp:simplePos x="0" y="0"/>
    <wp:positionH relativeFrom="page">
      <wp:posOffset>{_BORDER_POS_X}</wp:posOffset>
    </wp:positionH>
    <wp:positionV relativeFrom="page">
      <wp:posOffset>{_BORDER_POS_Y}</wp:posOffset>
    </wp:positionV>
    <wp:extent cx="{cx}" cy="{cy}"/>
    <wp:effectExtent l="0" t="0" r="0" b="0"/>
    <wp:wrapNone/>
    <wp:docPr id="{drawing_id}" name="PageBorderGroup{drawing_id}"/>
    <wp:cNvGraphicFramePr/>
    <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
      <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup">
        <wpg:wgp xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup">
          <wpg:cNvGrpSpPr/>
          <wpg:grpSpPr>
            <a:xfrm>
              <a:off x="0" y="0"/>
              <a:ext cx="{cx}" cy="{cy}"/>
              <a:chOff x="0" y="0"/>
              <a:chExt cx="{cx}" cy="{cy}"/>
            </a:xfrm>
          </wpg:grpSpPr>
          {shapes}
        </wpg:wgp>
      </a:graphicData>
    </a:graphic>
  </wp:anchor>
</w:drawing>"""
    return xml


# ---------------------------------------------------------------------------
# Logo
# ---------------------------------------------------------------------------
_BASE_DIR  = Path(__file__).parent.resolve()
_REPO_DIR = Path(__file__).resolve().parents[3]
_LOGO_CANDIDATES = [
    _REPO_DIR / "ONGC Logo.jpeg",
    _REPO_DIR / "ONGC Logo.jpg",
    _REPO_DIR / "ONGC Logo.png",
    _BASE_DIR / "ONGC Logo.jpeg",
    _BASE_DIR / "ONGC Logo.jpg",
    _BASE_DIR / "ONGC Logo.png",
    _BASE_DIR / "data" / "ongc_logo_for_docx.jpeg",
    _BASE_DIR / "data" / "ongc_logo_for_docx.jpg",
    _BASE_DIR / "data" / "ongc_logo_for_docx.png",
]
_LOGO_FALLBACK_PATH = _BASE_DIR / "data" / "ongc_logo_for_docx.png"


def _find_logo(logo_path: str | Path | None = None) -> Path | None:
    """Resolve ONGC logo from provided path, candidates, or embedded fallback."""
    if logo_path:
        p = Path(logo_path)
        if p.exists():
            return p
    for candidate in _LOGO_CANDIDATES:
        if candidate.exists():
            return candidate
    try:
        from embedded_logos import ONGC_LOGO_B64
        _LOGO_FALLBACK_PATH.parent.mkdir(parents=True, exist_ok=True)
        _LOGO_FALLBACK_PATH.write_bytes(base64.b64decode(ONGC_LOGO_B64))
        return _LOGO_FALLBACK_PATH
    except Exception as exc:
        logger.warning("Could not extract ONGC logo: %s", exc)
        return None


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def build_word_document(
    draft_dict: dict[str, Any],
    sections_dict: dict[str, str],
    parameters_list: list[dict[str, Any]],
    flags_list: list[dict[str, Any]] | None = None,
    impact_dict: dict[str, Any] | None = None,
    logo_path: str | Path | None = None,
) -> bytes:
    """Assemble ONGC spec sheet + CSC draft note sections as .docx bytes.

    Args:
        draft_dict: Dict with keys like spec_number, chemical_name, etc.
        sections_dict: Dict mapping section_name -> section_text
        parameters_list: List of parameter dicts with parameter_name, parameter_type, etc.
        flags_list: Optional issue flag rows
        impact_dict: Optional impact analysis dict
        logo_path: Optional path to logo file

    Returns:
        Document as bytes (BytesIO)
    """
    doc = Document()
    _configure_page(doc)
    _apply_base_styles(doc)
    _configure_spec_header(doc, logo_path)

    # ---- PART A: ONGC Specification Sheet -----------------------------------
    _build_spec_sheet(doc, draft_dict, parameters_list)

    # Page break before Part B
    doc.add_page_break()

    # ---- PART B: CSC Draft Note ---------------------------------------------
    _build_csc_draft_note(
        doc,
        draft_dict,
        sections_dict,
        parameters_list,
        flags_list or [],
        impact_dict,
        logo_path,
    )

    _add_footer(doc, draft_dict)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def build_flask_review_document(
    review_context: dict[str, Any],
    logo_path: str | Path | None = None,
) -> bytes:
    """Build a Flask-native workflow review document for one draft or revision."""
    doc = Document()
    _configure_page(doc)
    _apply_base_styles(doc)
    _configure_spec_header(doc, logo_path)

    spec_page_draft = _review_draft_as_spec_dict(review_context.get("draft") or {})
    spec_page_parameters = _review_parameters_as_spec_rows(review_context.get("parameter_rows", []) or [])
    _build_spec_sheet(doc, spec_page_draft, spec_page_parameters)

    doc.add_page_break()

    _build_review_cover_page(doc, review_context, logo_path)
    _build_review_snapshot_page(doc, review_context)
    _build_review_sections_page(doc, review_context)
    _build_review_parameters_page(doc, review_context)
    _build_review_appendix_page(doc, review_context)
    _add_footer(doc, review_context.get("draft") or {})

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _draft_value(draft: Any, field_name: str, default: Any = "—") -> Any:
    """Read a field from either a dict-like draft payload or an ORM object."""
    if isinstance(draft, dict):
        value = draft.get(field_name, default)
    else:
        value = getattr(draft, field_name, default)
    if value in (None, ""):
        return default
    return value


def _review_draft_as_spec_dict(draft: Any) -> dict[str, Any]:
    """Normalize review-context draft payloads into the Format A dict shape."""
    return {
        "spec_number": _draft_value(draft, "spec_number", "—"),
        "chemical_name": _draft_value(draft, "chemical_name", "—"),
        "material_code": _draft_value(draft, "material_code", "—"),
        "test_procedure": _draft_value(draft, "test_procedure", ""),
    }


def _review_parameters_as_spec_rows(parameter_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Build the Format A parameter rows from Flask review rows."""
    rows: list[dict[str, Any]] = []
    for parameter in parameter_rows:
        rows.append(
            {
                "parameter_name": parameter.get("parameter_name", "Untitled Parameter"),
                "parameter_type": parameter.get("parameter_type", ""),
                "existing_value": parameter.get("final_requirement", parameter.get("required_value", "—")),
                "proposed_value": "",
            }
        )
    return rows


def _build_review_cover_page(
    doc: Document,
    review_context: dict[str, Any],
    logo_path: str | Path | None = None,
) -> None:
    draft = review_context.get("draft") or {}
    logo_path = _find_logo(logo_path)

    header = doc.add_table(rows=1, cols=2)
    _set_table_no_borders(header)
    logo_cell, title_cell = header.rows[0].cells

    if logo_path and logo_path.exists():
        para = logo_cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.add_run().add_picture(str(logo_path), width=Inches(0.72))
    else:
        logo_cell.paragraphs[0].text = "ONGC"

    title_para = title_cell.paragraphs[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_para.add_run("CORPORATE SPECIFICATION REVIEW DOSSIER")
    title_run.bold = True
    title_run.font.name = _FONT_TITLE
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = _COPPER
    subtitle_para = title_cell.add_paragraph()
    subtitle_run = subtitle_para.add_run(
        "Workflow-generated approval package aligned with the Flask committee workbench."
    )
    subtitle_run.font.name = _FONT_MAIN
    subtitle_run.font.size = Pt(10)
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    _set_cell_widths(header.rows[0].cells, [0.95, 5.55])

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(str(_draft_value(draft, "spec_number")))
    run.bold = True
    run.font.name = _FONT_TITLE
    run.font.size = Pt(20)
    chemical = doc.add_paragraph()
    chemical.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chemical_run = chemical.add_run(str(_draft_value(draft, "chemical_name")))
    chemical_run.font.name = _FONT_MAIN
    chemical_run.font.size = Pt(15)
    chemical_run.bold = True
    chemical_run.font.color.rgb = _DARK_GRAY

    meta_table = doc.add_table(rows=0, cols=2)
    meta_table.style = "Table Grid"
    _set_table_outer_border(meta_table, _COPPER_HEX, sz=8)
    _set_table_inner_borders(meta_table, _LGRAY_HEX, sz=4)
    for row in review_context.get("summary_rows", [])[:8]:
        cells = meta_table.add_row().cells
        _shade_cell(cells[0], _LGRAY_HEX)
        label = cells[0].paragraphs[0].add_run(str(row.get("label", "Field")))
        label.bold = True
        label.font.name = _FONT_MAIN
        label.font.size = Pt(10)
        value = cells[1].paragraphs[0].add_run(str(row.get("value", "—")))
        value.font.name = _FONT_MAIN
        value.font.size = Pt(10)
        _set_cell_widths(cells, [2.2, 4.3])

    overview = doc.add_paragraph()
    overview.paragraph_format.space_before = Pt(8)
    overview_run = overview.add_run(
        "This dossier is a read-only submission snapshot intended for admin review, approval, return, or rejection."
    )
    overview_run.font.name = _FONT_MAIN
    overview_run.font.size = Pt(10)
    overview_run.font.color.rgb = _DARK_GRAY
    doc.add_page_break()


def _build_review_snapshot_page(doc: Document, review_context: dict[str, Any]) -> None:
    doc.add_heading("1. Submission Snapshot", level=1)
    summary_rows = review_context.get("summary_rows", []) or []
    if summary_rows:
        table = doc.add_table(rows=0, cols=4)
        table.style = "Table Grid"
        _set_table_outer_border(table, _COPPER_HEX, sz=8)
        _set_table_inner_borders(table, _LGRAY_HEX, sz=4)
        for index in range(0, len(summary_rows), 2):
            cells = table.add_row().cells
            pair = summary_rows[index:index + 2]
            for pair_index, row in enumerate(pair):
                label_cell = cells[pair_index * 2]
                value_cell = cells[pair_index * 2 + 1]
                _shade_cell(label_cell, _LGRAY_HEX)
                label = label_cell.paragraphs[0].add_run(str(row.get("label", "Field")))
                label.bold = True
                label.font.name = _FONT_MAIN
                label.font.size = Pt(9)
                value = value_cell.paragraphs[0].add_run(str(row.get("value", "—")))
                value.font.name = _FONT_MAIN
                value.font.size = Pt(9)
            _set_cell_widths(cells, [1.35, 1.95, 1.35, 1.95])

    doc.add_page_break()


def _build_review_sections_page(doc: Document, review_context: dict[str, Any]) -> None:
    doc.add_heading("2. Section Narrative", level=1)
    for section in review_context.get("section_rows", []) or []:
        doc.add_heading(str(section.get("label", "Section")), level=2)
        paragraph = doc.add_paragraph(str(section.get("text", "") or "—"))
        paragraph.style.font.name = _FONT_MAIN

    for title, rows in (
        ("3. Material Master Snapshot", review_context.get("master_rows", [])),
        ("4. Material Properties", review_context.get("material_property_rows", [])),
        ("5. Storage and Handling", review_context.get("storage_rows", [])),
        ("6. Impact Assessment", review_context.get("impact_rows", [])),
    ):
        doc.add_heading(title, level=1)
        _add_key_value_table(doc, rows)

    doc.add_page_break()


def _build_review_parameters_page(doc: Document, review_context: dict[str, Any]) -> None:
    doc.add_heading("7. Parameter Review and Decision Basis", level=1)
    parameter_rows = review_context.get("parameter_rows", []) or []
    if not parameter_rows:
        doc.add_paragraph("No parameters are present in this submission.")
        return

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    _set_table_outer_border(table, _COPPER_HEX, sz=8)
    _set_table_inner_borders(table, _LGRAY_HEX, sz=4)
    headers = [
        "Parameter",
        "Type",
        "Unit",
        "Current Requirement",
        "Proposed Requirement",
        "Change Status",
    ]
    widths = [1.5, 0.7, 0.7, 1.45, 1.45, 0.8]
    for idx, header in enumerate(headers):
        _shade_cell(table.rows[0].cells[idx], _COPPER_HEX, font_white=True)
        paragraph = table.rows[0].cells[idx].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        run.bold = True
        run.font.name = _FONT_MAIN
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _set_cell_widths(table.rows[0].cells, widths)

    for parameter in parameter_rows:
        cells = table.add_row().cells
        values = [
            str(parameter.get("parameter_name", "Untitled Parameter")),
            str(parameter.get("parameter_type", "—")),
            str(parameter.get("unit_of_measure", "—")),
            str(parameter.get("required_value", "—")),
            str(parameter.get("proposed_value", "No change submitted")),
            str(parameter.get("change_status", "Retained")),
        ]
        for idx, value in enumerate(values):
            paragraph = cells[idx].paragraphs[0]
            run = paragraph.add_run(value)
            run.font.name = _FONT_MAIN
            run.font.size = Pt(9)
            if idx == 4 and parameter.get("change_status") == "Revised":
                run.bold = True
                run.font.color.rgb = _COPPER
            if idx == 5:
                run.bold = True
                run.font.color.rgb = _COPPER if parameter.get("change_status") == "Revised" else _DARK_GRAY
        _set_cell_widths(cells, widths)

    doc.add_heading("8. Parameter Technical Fields", level=1)
    support = doc.add_table(rows=1, cols=4)
    support.style = "Table Grid"
    _set_table_outer_border(support, _COPPER_HEX, sz=8)
    _set_table_inner_borders(support, _LGRAY_HEX, sz=4)
    support_headers = ["Parameter", "Conditions", "Test Procedure", "Procedure Text"]
    for idx, header in enumerate(support_headers):
        _shade_cell(support.rows[0].cells[idx], _LGRAY_HEX)
        run = support.rows[0].cells[idx].paragraphs[0].add_run(header)
        run.bold = True
        run.font.name = _FONT_MAIN
        run.font.size = Pt(9)
    _set_cell_widths(support.rows[0].cells, [1.7, 1.6, 1.2, 2.0])
    for parameter in parameter_rows:
        cells = support.add_row().cells
        values = [
            str(parameter.get("parameter_name", "Untitled Parameter")),
            str(parameter.get("conditions", "—")),
            str(parameter.get("test_procedure_type", "—")),
            str(parameter.get("procedure_text", "—")),
        ]
        for idx, value in enumerate(values):
            run = cells[idx].paragraphs[0].add_run(value)
            run.font.name = _FONT_MAIN
            run.font.size = Pt(9)
        _set_cell_widths(cells, [1.7, 1.6, 1.2, 2.0])

    doc.add_page_break()


def _build_review_appendix_page(doc: Document, review_context: dict[str, Any]) -> None:
    doc.add_heading("9. Review Notes and Audit Context", level=1)
    revision = review_context.get("revision")
    reviewer_notes = getattr(revision, "reviewer_notes", "") if revision is not None else ""
    if reviewer_notes:
        doc.add_paragraph(str(reviewer_notes))
    else:
        doc.add_paragraph("No prior admin review note has been recorded for this submission.")


def _add_key_value_table(doc: Document, rows: list[dict[str, Any]] | None) -> None:
    rows = rows or []
    if not rows:
        doc.add_paragraph("No data captured for this section.")
        return

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    _set_table_outer_border(table, _COPPER_HEX, sz=8)
    _set_table_inner_borders(table, _LGRAY_HEX, sz=4)
    for row in rows:
        cells = table.add_row().cells
        _shade_cell(cells[0], _LGRAY_HEX)
        label = cells[0].paragraphs[0].add_run(str(row.get("label", "Field")))
        label.bold = True
        label.font.name = _FONT_MAIN
        label.font.size = Pt(9)
        value = cells[1].paragraphs[0].add_run(str(row.get("value", "—")))
        value.font.name = _FONT_MAIN
        value.font.size = Pt(9)
        _set_cell_widths(cells, [2.0, 4.5])


def build_master_spec_document(
    all_specs: list[dict[str, Any]],
    include_draft_note: bool = False,
    changelog: list[dict[str, Any]] | None = None,
    include_metadata: bool = False,
    logo_path: str | Path | None = None,
) -> bytes:
    """Build a single Word document containing multiple ONGC spec sheets.

    Each item in all_specs must have keys:
        draft      – dict with spec_number, chemical_name, etc.
        sections   – dict mapping section_name -> section_text
        parameters – list of parameter dicts
        flags      – list of issue flags
        impact_analysis – optional impact analysis dict

    Each spec gets its own ONGC spec sheet page (Part A).
    If include_draft_note=True, a CSC Draft Note section (Part B) follows each spec.
    If changelog is provided, a final Change Log page is appended after all specs.
    Specs are separated by page breaks.
    """
    doc = Document()
    _configure_page(doc)
    _apply_base_styles(doc)
    _configure_spec_header(doc, logo_path)

    ordered_specs = sorted(
        all_specs,
        key=lambda item: spec_sort_key(
            str(item.get("draft", {}).get("spec_number", "") or "")
        ),
    )

    if ordered_specs:
        _build_master_index_page(
            doc,
            ordered_specs,
            include_draft_note=include_draft_note,
            logo_path=logo_path,
        )
        doc.add_page_break()

    for i, item in enumerate(ordered_specs):
        draft      = item["draft"]
        sections   = item.get("sections", {})
        parameters = item.get("parameters", [])
        flags      = item.get("flags", [])

        if isinstance(sections, dict):
            sections_dict = {
                str(key): str(value or "")
                for key, value in sections.items()
            }
        else:
            sections_dict = {
                s.get("section_name", ""): s.get("section_text", "")
                for s in (sections or [])
            }

        if i > 0:
            doc.add_page_break()

        # Part A — ONGC Spec Sheet
        _build_spec_sheet(doc, draft, parameters)

        if include_draft_note:
            doc.add_page_break()
            _build_csc_draft_note(
                doc, draft, sections_dict, parameters, flags,
                item.get("impact_analysis"), logo_path
            )

    # Final page — Change Log
    if changelog is not None:
        doc.add_page_break()
        _build_change_log_page(doc, changelog, logo_path)

    if include_metadata and ordered_specs:
        doc.add_page_break()
        _build_metadata_appendix_page(doc, ordered_specs, logo_path)

    _add_master_footer(doc, len(ordered_specs))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ===========================================================================
# PART A — ONGC Spec Sheet (mirrors CSC Format A layout exactly)
# ===========================================================================

def _build_spec_sheet(
    doc: Document,
    draft: dict[str, Any],
    parameters: list[dict[str, Any]],
) -> None:
    """Build the ONGC Corporate Specification sheet page."""

    # ── Double-line page border (DrawingML, behind text) ─────────────────
    border_anchor = doc.add_paragraph()
    _add_page_border(doc, border_anchor)

    # ── Header info table (4 rows × 2 cols) ──────────────────────────────
    _add_spec_info_table(doc, draft)

    # ── Parameter table (4-col: S.No | Parameter(Type) | : | Value) ──────
    _add_spec_parameters(doc, parameters)


def _add_page_border(doc: Document, anchor_paragraph=None) -> None:
    """Insert the DrawingML border into the provided paragraph."""
    para = anchor_paragraph if anchor_paragraph is not None else doc.add_paragraph()
    border_xml = _make_border_drawing_xml(_next_drawing_id())
    run = para.add_run()
    run._r.append(parse_xml(border_xml))


def _configure_spec_header(doc: Document, logo_path: str | Path | None = None) -> None:
    """Configure a reusable document header for all spec pages."""
    logo_path = _find_logo(logo_path)
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    # Reset current header content to avoid carrying stale elements.
    for para in list(header.paragraphs):
        para._element.getparent().remove(para._element)
    for table in list(header.tables):
        table._element.getparent().remove(table._element)

    tbl = header.add_table(rows=1, cols=2, width=Twips(10_031))
    _set_table_no_borders(tbl)
    _set_table_width_dxa(tbl, 10_031)
    _set_table_grid_dxa(tbl, [1_617, 8_414])

    logo_cell  = tbl.rows[0].cells[0]
    title_cell = tbl.rows[0].cells[1]

    # Logo
    if logo_path and logo_path.exists():
        lp = logo_cell.paragraphs[0]
        lp.paragraph_format.left_indent = Twips(-142)
        r = lp.add_run()
        r.add_picture(str(logo_path), width=Emu(980_118), height=Emu(432_630))
    else:
        logo_cell.paragraphs[0].text = "ONGC"

    # Title
    tp = title_cell.paragraphs[0]
    tp.paragraph_format.space_before = Pt(4)
    tp.paragraph_format.right_indent = Twips(-848)
    r = tp.add_run("CORPORATE SPECIFICATIONS FOR OIL FIELD CHEMICALS")
    r.bold = True
    r.font.name = _FONT_TITLE
    r.font.size = Pt(14)
    r.font.color.rgb = _BLACK

    _set_cell_widths_dxa(tbl.rows[0].cells, [1_617, 8_414])
    sp = header.add_paragraph()
    sp.style = doc.styles["Header"]
    sp.paragraph_format.space_after = Pt(2)


def _add_spec_info_table(doc: Document, draft: dict[str, Any]) -> None:
    """4-row × 2-col header info table exactly matching Format A.

    Col widths: 2.58" label | 4.02" value
    Borders: single, sz=4 (0.5pt), black — applied per-cell like Format A.
    No shading/fill.
    """
    tbl = doc.add_table(rows=4, cols=2)
    # Use TableGrid as base (gives all cells borders) then we'll refine
    tbl.style = "Table Grid"
    _set_table_indent_dxa(tbl, 392)
    _set_table_grid_dxa(tbl, [3_323, 5_789])

    rows_data = [
        ("CHEMICAL NAME",
         (draft.get("chemical_name", "") or "—").upper(),
         True,  True),
        ("SPECIFICATION No.",
         draft.get("spec_number", "") or "—",
         True,  True),
        ("TEST PROCEDURE",
         draft.get("test_procedure", "") or
         (draft.get("spec_number", "").replace("/", "/ TP /") if draft.get("spec_number") else "—"),
         True,  True),
        ("MATERIAL CODE",
         draft.get("material_code", "") or "—",
         True,  False),
    ]

    for row_idx, (label, value, label_bold, value_bold) in enumerate(rows_data):
        cells = tbl.rows[row_idx].cells

        # Label cell
        lp = cells[0].paragraphs[0]
        lr = lp.add_run(label)
        lr.bold      = label_bold
        lr.font.name = _FONT_MAIN
        lr.font.size = Pt(13)
        lr.font.color.rgb = _BLACK

        # Value cell
        vp = cells[1].paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr = vp.add_run(value)
        vr.bold      = value_bold
        vr.font.name = _FONT_MAIN
        vr.font.size = Pt(13)
        vr.font.color.rgb = _BLACK

    # Set col widths exactly to cleaned layout (dxa twips)
    for row in tbl.rows:
        _set_cell_widths_dxa(row.cells, [3_323, 5_789])

    # Override borders to match Format A: single 0.5pt black (sz=4)
    _set_table_format_a_borders(tbl)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(4)
    sp.paragraph_format.space_after  = Pt(4)


def _add_spec_parameters(doc: Document, parameters: list[dict[str, Any]]) -> None:
    """Parameter table: 4 cols — S.No | Parameter Name (Type) | : | Required Value

    Matches Format A Table 3:
      Col 0 (S.No.)     : 0.56"  (805 dxa)
      Col 1 (Parameter) : 2.72"  (3910 dxa) — name + (Type) in brackets
      Col 2 (:)         : 0.20"  (290 dxa)
      Col 3 (Value)     : 3.01"  (4340 dxa)
    TableGrid style, Times New Roman 11pt, no shading.
    """
    if not parameters:
        p = doc.add_paragraph("No parameters recorded.")
        p.paragraph_format.space_before = Pt(4)
        return

    # ── Detect group headers (parameter_name starts with [GroupName] ...) ──
    grouped: list[tuple[str | None, list[dict]]] = []
    current_group: str | None = None
    current_items: list[dict]  = []

    for param in parameters:
        name  = param.get("parameter_name", "")
        grp_m = re.match(r"^\[(.+?)\]\s*(.*)", name)
        if grp_m:
            g          = grp_m.group(1)
            clean_name = grp_m.group(2)
            if g != current_group:
                if current_items:
                    grouped.append((current_group, current_items))
                current_group  = g
                current_items  = []
            current_items.append({**param, "parameter_name": clean_name})
        else:
            current_items.append(param)

    if current_items:
        grouped.append((current_group, current_items))
    if not grouped:
        grouped = [(None, parameters)]

    param_counter = 1

    for group_label, items in grouped:
        # ── Group header row ──
        if group_label:
            gh = doc.add_paragraph()
            gh.paragraph_format.left_indent  = Inches(0.15)
            gh.paragraph_format.space_before = Pt(6)
            gh.paragraph_format.space_after  = Pt(2)
            gr = gh.add_run(f"{group_label}:")
            gr.bold       = True
            gr.underline  = True
            gr.font.name  = _FONT_MAIN
            gr.font.size  = Pt(11)
            gr.font.color.rgb = _BLACK

        # ── 4-column parameter table for this group ──
        tbl = doc.add_table(rows=len(items) + 1, cols=4)
        tbl.style = "Table Grid"
        _set_table_indent_dxa(tbl, 392)
        _set_table_grid_dxa(tbl, [540, 3917, 290, 4334])

        # Header row
        hdr = tbl.rows[0].cells
        for ci, hdr_text in enumerate(["S. No.", "Parameter", "", "Required Value"]):
            hp = hdr[ci].paragraphs[0]
            if ci == 2:
                hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hr_run = hp.add_run(hdr_text)
            hr_run.bold       = True
            hr_run.font.name  = _FONT_MAIN
            hr_run.font.size  = Pt(11)
            hr_run.font.color.rgb = _BLACK

        # Data rows
        for row_idx, param in enumerate(items, start=1):
            cells = tbl.rows[row_idx].cells
            name     = param.get("parameter_name", "")
            ptype    = param.get("parameter_type", "")
            proposed = (
                param.get("proposed_value") or
                param.get("existing_value") or ""
            )
            existing = param.get("existing_value", "") or ""

            # Col 0 — S. No.
            np_ = cells[0].paragraphs[0]
            np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
            nr = np_.add_run(str(param_counter))
            nr.font.name  = _FONT_MAIN
            nr.font.size  = Pt(11)
            nr.font.color.rgb = _BLACK

            # Col 1 — Parameter Name (Type)
            # Append (type) in brackets if type is set
            type_suffix = f"  ({ptype})" if ptype else ""
            nm_p = cells[1].paragraphs[0]
            nm_p.paragraph_format.left_indent = Inches(0.05)
            nm_run = nm_p.add_run(name)
            nm_run.bold       = True
            nm_run.font.name  = _FONT_MAIN
            nm_run.font.size  = Pt(11)
            nm_run.font.color.rgb = _BLACK
            if type_suffix:
                ts_run = nm_p.add_run(type_suffix)
                ts_run.bold      = False
                ts_run.italic    = True
                ts_run.font.name = _FONT_MAIN
                ts_run.font.size = Pt(10)
                ts_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

            # Col 2 — Colon separator
            cp_ = cells[2].paragraphs[0]
            cp_.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cp_.add_run(":")
            cr.bold       = True
            cr.font.name  = _FONT_MAIN
            cr.font.size  = Pt(11)
            cr.font.color.rgb = _BLACK

            # Col 3 — Required Value
            vp_ = cells[3].paragraphs[0]
            vp_.paragraph_format.left_indent = Inches(0.05)
            vr = vp_.add_run(proposed)
            vr.bold   = True
            vr.italic = True
            vr.font.name = _FONT_MAIN
            vr.font.size = Pt(11)
            # Highlight proposed changes in copper
            if proposed.strip() and existing.strip() and proposed.strip() != existing.strip():
                vr.font.color.rgb = _COPPER
            else:
                vr.font.color.rgb = _BLACK

            param_counter += 1

        # Apply cleaned-layout cell widths to every row (dxa twips)
        for row in tbl.rows:
            _set_cell_widths_dxa(row.cells, [414, 3917, 290, 4334])

    # Bottom spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)


# ===========================================================================
# MASTER INDEX PAGE
# ===========================================================================

def _build_master_index_page(
    doc: Document,
    all_specs: list[dict[str, Any]],
    include_draft_note: bool = False,
    logo_path: str | Path | None = None,
) -> None:
    """Build index page with spec order and starting page numbers."""
    logo_path = _find_logo(logo_path)

    hdr_tbl = doc.add_table(rows=1, cols=2)
    _set_table_no_borders(hdr_tbl)

    logo_cell = hdr_tbl.rows[0].cells[0]
    title_cell = hdr_tbl.rows[0].cells[1]

    if logo_path and logo_path.exists():
        lp = logo_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lp.add_run().add_picture(str(logo_path), width=Inches(0.55))
    else:
        logo_cell.paragraphs[0].text = "ONGC"

    tp = title_cell.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tp.paragraph_format.space_before = Pt(4)
    tr = tp.add_run("CORPORATE SPECIFICATIONS — MASTER INDEX")
    tr.bold = True
    tr.font.name = _FONT_TITLE
    tr.font.size = Pt(14)
    tr.font.color.rgb = _BLACK
    _set_cell_widths(hdr_tbl.rows[0].cells, [0.65, 5.85])

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)

    subset_line = ", ".join(SPEC_SUBSET_ORDER)
    subset_desc = "; ".join(
        f"{code}: {SPEC_SUBSET_LABELS.get(code, code)}" for code in SPEC_SUBSET_ORDER
    )

    order_note = doc.add_paragraph(
        f"Order: {subset_line}"
    )
    order_note.paragraph_format.space_after = Pt(1)
    nr = order_note.runs[0]
    nr.font.name = _FONT_MAIN
    nr.font.size = Pt(9)
    nr.italic = True
    nr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    order_desc = doc.add_paragraph(subset_desc)
    order_desc.paragraph_format.space_after = Pt(6)
    dr = order_desc.runs[0]
    dr.font.name = _FONT_MAIN
    dr.font.size = Pt(8)
    dr.italic = True
    dr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    col_headers = ["S. No.", "Spec No.", "Name of Chemical", "Material Code", "Page No."]
    col_widths = [0.45, 1.65, 2.35, 1.35, 0.65]

    tbl = doc.add_table(rows=1, cols=len(col_headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_format_a_borders(tbl)

    hdr_cells = tbl.rows[0].cells
    for ci, hdr_text in enumerate(col_headers):
        _shade_cell(hdr_cells[ci], "1F2937")
        hp = hdr_cells[ci].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(hdr_text)
        hr.bold = True
        hr.font.name = _FONT_MAIN
        hr.font.size = Pt(10)
        hr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _set_cell_widths(hdr_cells, col_widths)

    pages_per_spec = 2 if include_draft_note else 1
    first_spec_page = 2

    for idx, item in enumerate(all_specs, start=1):
        draft = item.get("draft", {}) or {}
        spec_no = str(draft.get("spec_number", "") or "—")
        chem_name = str(draft.get("chemical_name", "") or "—")
        material_code = str(draft.get("material_code", "") or "—")
        subset_code, _, _ = parse_spec_number(spec_no)
        if subset_code in SPEC_SUBSET_LABELS:
            chem_name = f"{chem_name} [{subset_code}]"

        row_cells = tbl.add_row().cells
        row_data = [
            str(idx),
            spec_no,
            chem_name,
            material_code,
            str(first_spec_page + (idx - 1) * pages_per_spec),
        ]
        for ci, value in enumerate(row_data):
            rp = row_cells[ci].paragraphs[0]
            if ci in (0, 4):
                rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = rp.add_run(value)
            rr.font.name = _FONT_MAIN
            rr.font.size = Pt(10)
            rr.font.color.rgb = _BLACK
        _set_cell_widths(row_cells, col_widths)

    sp2 = doc.add_paragraph()
    sp2.paragraph_format.space_after = Pt(6)


# ===========================================================================
# PART B — CSC Draft Note sections
# ===========================================================================

def _build_csc_draft_note(
    doc: Document,
    draft: dict[str, Any],
    section_map: dict[str, str],
    parameters: list[dict[str, Any]],
    flags: list[dict[str, Any]],
    impact_analysis: dict[str, Any] | None = None,
    logo_path: str | Path | None = None,
) -> None:
    """Build the committee note section (Background, Issues, Impact, Recommendation)."""
    _add_draft_note_header(doc, draft, logo_path)
    _add_revision_summary_section(doc, draft, section_map, parameters, flags, impact_analysis)

    _add_note_section(doc, "1. Background Context",
                      section_map.get(SECTION_BACKGROUND, ""))
    _add_note_section(doc, "2. Existing Specification Summary",
                      section_map.get(SECTION_EXISTING_SPEC, ""))
    _add_issues_section(doc, flags)
    _add_parameter_review_section(doc, parameters)
    _add_note_section(doc, "5. Proposed Changes Narrative",
                      section_map.get(SECTION_PROPOSED, ""))
    _add_note_section(doc, "6. Technical / Operational Justification",
                      section_map.get(SECTION_JUSTIFICATION, ""))
    _add_impact_section(doc, impact_analysis, section_map)
    _add_recommendation_section(doc, section_map)


def _add_revision_summary_section(
    doc: Document,
    draft: dict[str, Any],
    section_map: dict[str, str],
    parameters: list[dict[str, Any]],
    flags: list[dict[str, Any]],
    impact_analysis: dict[str, Any] | None = None,
) -> None:
    """Add a structured summary before the detailed committee note."""
    doc.add_heading("Executive Summary of Proposal for Revision of Corporate Specification", level=1)

    total_parameters = len(parameters or [])
    changed_parameters = _collect_changed_parameters(parameters)
    vital_count = sum(
        1 for parameter in (parameters or [])
        if str(parameter.get("parameter_type", "") or "").strip().lower() == "vital"
    )
    desirable_count = total_parameters - vital_count
    flagged_count = sum(1 for flag in (flags or []) if bool(flag.get("is_present")))
    version_label = str(
        draft.get("version_display")
        or draft.get("version")
        or f"v{draft.get('spec_version', 0)}"
    )
    impact_grade = "—"
    if impact_analysis:
        impact_grade = str(impact_analysis.get("impact_grade") or "—")

    overview = doc.add_table(rows=3, cols=4)
    overview.style = "Table Grid"
    _set_table_outer_border(overview, _COPPER_HEX, sz=8)
    _set_table_inner_borders(overview, _LGRAY_HEX, sz=4)

    summary_pairs = [
        ("Specification No.", draft.get("spec_number", "—")),
        ("Version", version_label),
        ("Chemical Name", draft.get("chemical_name", "—")),
        ("Material Code", draft.get("material_code", "—")),
        ("Total Parameters", str(total_parameters)),
        ("Changed Parameters", str(len(changed_parameters))),
        ("Vital Parameters", str(vital_count)),
        ("Desirable Parameters", str(desirable_count)),
        ("Flagged Issues", str(flagged_count)),
        ("Impact Grade", impact_grade),
        ("Status", draft.get("status", "—")),
        ("Reviewed By", draft.get("reviewed_by", "—")),
    ]
    for row_idx, row in enumerate(overview.rows):
        pairs = summary_pairs[row_idx * 2:(row_idx + 1) * 2]
        cells = row.cells
        for pair_idx, (label, value) in enumerate(pairs):
            label_cell = cells[pair_idx * 2]
            value_cell = cells[pair_idx * 2 + 1]
            _shade_cell(label_cell, _LGRAY_HEX)
            label_paragraph = label_cell.paragraphs[0]
            label_run = label_paragraph.add_run(label)
            label_run.bold = True
            label_run.font.name = _FONT_MAIN
            label_run.font.size = Pt(10)
            value_paragraph = value_cell.paragraphs[0]
            value_run = value_paragraph.add_run(str(value or "—"))
            value_run.font.name = _FONT_MAIN
            value_run.font.size = Pt(10)
    for row in overview.rows:
        _set_cell_widths(row.cells, [1.4, 1.8, 1.4, 1.8])

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    _add_summary_table(
        doc,
        [
            ("Background Context", section_map.get(SECTION_BACKGROUND, "")),
            ("Existing Specification Summary", section_map.get(SECTION_EXISTING_SPEC, "")),
            ("Proposed Changes", section_map.get(SECTION_PROPOSED, "")),
            ("Technical / Operational Justification", section_map.get(SECTION_JUSTIFICATION, "")),
            ("Version Change Reason", section_map.get(REC_MAIN_KEY, "")),
            ("Additional Remarks", section_map.get(REC_REMARKS_KEY, "")),
        ],
    )

    doc.add_heading("Key Parameter Revisions", level=2)
    if changed_parameters:
        _add_parameter_change_table(doc, changed_parameters)
    else:
        paragraph = doc.add_paragraph(
            "No explicit proposed value changes were captured. The export still includes the current specification sheet and detailed draft note."
        )
        paragraph.style.font.name = _FONT_MAIN
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _add_summary_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    _set_table_outer_border(table, _COPPER_HEX, sz=8)
    _set_table_inner_borders(table, _LGRAY_HEX, sz=4)

    for label, value in rows:
        row = table.add_row().cells
        _shade_cell(row[0], _LGRAY_HEX)
        label_paragraph = row[0].paragraphs[0]
        label_run = label_paragraph.add_run(label)
        label_run.bold = True
        label_run.font.name = _FONT_MAIN
        label_run.font.size = Pt(10)
        value_paragraph = row[1].paragraphs[0]
        value_run = value_paragraph.add_run((value or "—").strip() or "—")
        value_run.font.name = _FONT_MAIN
        value_run.font.size = Pt(10)
        _set_cell_widths(row, [2.1, 4.4])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _collect_changed_parameters(parameters: list[dict[str, Any]] | None) -> list[dict[str, Any]]:
    changed: list[dict[str, Any]] = []
    for index, parameter in enumerate(parameters or [], start=1):
        existing = str(parameter.get("existing_value", "") or "").strip()
        proposed = str(parameter.get("proposed_value", "") or "").strip()
        if not proposed or proposed == existing:
            continue
        changed.append(
            {
                "index": index,
                "parameter_name": str(parameter.get("parameter_name", "") or "—"),
                "parameter_type": str(parameter.get("parameter_type", "") or "—"),
                "existing_value": existing or "—",
                "proposed_value": proposed or "—",
            }
        )
    return changed


def _add_parameter_change_table(doc: Document, changed_parameters: list[dict[str, Any]]) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    _set_table_outer_border(table, _COPPER_HEX, sz=8)
    _set_table_inner_borders(table, _LGRAY_HEX, sz=4)

    headers = ["S. No.", "Parameter", "Type", "Existing Requirement", "Proposed Requirement"]
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        _shade_cell(header_cells[idx], _COPPER_HEX, font_white=True)
        paragraph = header_cells[idx].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        run.bold = True
        run.font.name = _FONT_MAIN
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _set_cell_widths(header_cells, [0.55, 2.0, 0.9, 1.7, 1.7])

    for item in changed_parameters:
        row = table.add_row().cells
        values = [
            str(item["index"]),
            item["parameter_name"],
            item["parameter_type"],
            item["existing_value"],
            item["proposed_value"],
        ]
        for idx, value in enumerate(values):
            paragraph = row[idx].paragraphs[0]
            if idx == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(value)
            run.font.name = _FONT_MAIN
            run.font.size = Pt(10)
            if idx == 4:
                run.bold = True
                run.font.color.rgb = _COPPER
        _set_cell_widths(row, [0.55, 2.0, 0.9, 1.7, 1.7])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _build_metadata_appendix_page(
    doc: Document,
    ordered_specs: list[dict[str, Any]],
    logo_path: str | Path | None = None,
) -> None:
    """Append a compact metadata appendix for the exported master set."""
    logo_path = _find_logo(logo_path)

    header_table = doc.add_table(rows=1, cols=2)
    _set_table_no_borders(header_table)

    logo_cell = header_table.rows[0].cells[0]
    title_cell = header_table.rows[0].cells[1]

    if logo_path and logo_path.exists():
        logo_paragraph = logo_cell.paragraphs[0]
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_paragraph.add_run().add_picture(str(logo_path), width=Inches(0.55))
    else:
        logo_cell.paragraphs[0].text = "ONGC"

    title_paragraph = title_cell.paragraphs[0]
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_paragraph.paragraph_format.space_before = Pt(4)
    title_run = title_paragraph.add_run("CORPORATE SPECIFICATIONS — METADATA APPENDIX")
    title_run.bold = True
    title_run.font.name = _FONT_TITLE
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = _BLACK
    _set_cell_widths(header_table.rows[0].cells, [0.65, 5.85])

    intro = doc.add_paragraph(
        "Appendix generated from the current published specification records included in this export."
    )
    intro.paragraph_format.space_after = Pt(6)
    intro.runs[0].font.name = _FONT_MAIN
    intro.runs[0].font.size = Pt(10)
    intro.runs[0].italic = True
    intro.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    _set_table_format_a_borders(table)

    headers = [
        "S. No.",
        "Spec No.",
        "Chemical",
        "Subset",
        "Version",
        "Parameters",
        "Meeting Date",
        "Reviewed By",
    ]
    widths = [0.4, 1.2, 1.8, 0.8, 0.6, 0.7, 0.9, 1.1]

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        _shade_cell(header_cells[index], "1F2937")
        paragraph = header_cells[index].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        run.bold = True
        run.font.name = _FONT_MAIN
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _set_cell_widths(header_cells, widths)

    for index, item in enumerate(ordered_specs, start=1):
        draft = item.get("draft", {}) or {}
        subset_code, _, _ = parse_spec_number(str(draft.get("spec_number", "") or ""))
        row = table.add_row().cells
        values = [
            str(index),
            str(draft.get("spec_number", "") or "—"),
            str(draft.get("chemical_name", "") or "—"),
            subset_code or "—",
            str(draft.get("version_display") or draft.get("version") or f"v{draft.get('spec_version', 0)}"),
            str(len(item.get("parameters", []) or [])),
            str(draft.get("meeting_date", "") or "—"),
            str(draft.get("reviewed_by", "") or "—"),
        ]
        for cell_index, value in enumerate(values):
            paragraph = row[cell_index].paragraphs[0]
            if cell_index in {0, 4, 5}:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(value)
            run.font.name = _FONT_MAIN
            run.font.size = Pt(9)
            run.font.color.rgb = _BLACK
        _set_cell_widths(row, widths)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def _add_draft_note_header(doc: Document, draft: dict[str, Any], logo_path: str | Path | None = None) -> None:
    """Formal CSC export cover block aligned with the Flask workflow review language."""
    logo_path = _find_logo(logo_path)

    tbl = doc.add_table(rows=1, cols=2)
    _set_table_no_borders(tbl)

    logo_cell  = tbl.rows[0].cells[0]
    title_cell = tbl.rows[0].cells[1]

    if logo_path and logo_path.exists():
        lp = logo_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lp.add_run().add_picture(str(logo_path), width=Inches(0.55))
    else:
        logo_cell.paragraphs[0].text = "ONGC"

    tp = title_cell.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = tp.add_run("CORPORATE SPECIFICATION REVISION MEMORANDUM")
    r.bold = True
    r.font.name = _FONT_TITLE
    r.font.size = Pt(13)
    r.font.color.rgb = _COPPER

    subtitle = title_cell.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle.paragraph_format.space_before = Pt(2)
    subtitle_run = subtitle.add_run(
        "Workflow-generated review package. Read-only submission snapshot for approval, return, or archival reference."
    )
    subtitle_run.font.name = _FONT_MAIN
    subtitle_run.font.size = Pt(9)
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    _set_cell_widths(tbl.rows[0].cells, [0.65, 5.85])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Metadata block
    meta = doc.add_table(rows=3, cols=4)
    meta.style = "Table Grid"
    _set_table_outer_border(meta, _COPPER_HEX, sz=8)
    _set_table_inner_borders(meta, _LGRAY_HEX, sz=4)

    meta_data = [
        ("Specification No.", draft.get("spec_number", "—"),
         "Chemical Name",     draft.get("chemical_name", "—")),
        ("Prepared By",       draft.get("prepared_by", "—"),
         "Reviewed By",       draft.get("reviewed_by", "—")),
        ("Meeting Date",      draft.get("meeting_date", "—"),
         "Status",            draft.get("status", "Draft")),
    ]
    for ri, (l1, v1, l2, v2) in enumerate(meta_data):
        cells = meta.rows[ri].cells
        for ci, (lbl, val) in enumerate([(l1, v1), (l2, v2)]):
            label_cell = cells[ci * 2]
            value_cell = cells[ci * 2 + 1]
            _shade_cell(label_cell, _LGRAY_HEX)
            lp2 = label_cell.paragraphs[0]
            lr2 = lp2.add_run(lbl)
            lr2.bold = True
            lr2.font.name = _FONT_MAIN
            lr2.font.size = Pt(10)
            vp2 = value_cell.paragraphs[0]
            vr2 = vp2.add_run(val)
            vr2.font.name = _FONT_MAIN
            vr2.font.size = Pt(10)

    for row in meta.rows:
        _set_cell_widths(row.cells, [1.4, 2.0, 1.4, 1.7])

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _add_note_section(doc: Document, title: str, text: str) -> None:
    doc.add_heading(title, level=1)
    body = text.strip() or "—"
    p = doc.add_paragraph(body)
    p.style.font.name = _FONT_MAIN
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _add_issues_section(doc: Document, flags: list[dict[str, Any]]) -> None:
    doc.add_heading("3. Issue Flag Register", level=1)
    normalized_flags = flags or []
    flag_map = {str(f.get("issue_type", "")): f for f in normalized_flags}

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    _set_table_outer_border(tbl, "000000", sz=6)
    _set_table_inner_borders(tbl, _LGRAY_HEX, sz=4)

    hdr = tbl.rows[0].cells
    for i, col in enumerate(["Issue Type", "Present", "Details / Note"]):
        _shade_cell(hdr[i], _COPPER_HEX, font_white=True)
        hp = hdr[i].paragraphs[0]
        hr = hp.add_run(col)
        hr.bold = True
        hr.font.name = _FONT_MAIN
        hr.font.size = Pt(10)
        hr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    issue_rows: list[tuple[str, bool, str]] = []
    for key, label in ISSUE_TYPES:
        flag = flag_map.get(key)
        if flag is None:
            flag = next(
                (
                    item for item in normalized_flags
                    if str(item.get("issue_type", "")).strip().lower() == str(label).strip().lower()
                ),
                None,
            )
        issue_rows.append(
            (
                label,
                bool(flag.get("is_present", 0)) if flag else False,
                str(flag.get("note", "") or "") if flag else "",
            )
        )

    extra_labels = [
        str(item.get("issue_type", "") or "").strip()
        for item in normalized_flags
        if str(item.get("issue_type", "") or "").strip()
        and str(item.get("issue_type", "") or "").strip() not in {key for key, _ in ISSUE_TYPES}
        and str(item.get("issue_type", "") or "").strip().lower() not in {label.lower() for _, label in ISSUE_TYPES}
    ]
    for label in extra_labels:
        flag = flag_map.get(label, {})
        issue_rows.append((label, bool(flag.get("is_present", 0)), str(flag.get("note", "") or "")))

    for label, is_yes, note in issue_rows:
        row    = tbl.add_row().cells
        for ci, val in enumerate([label, "Yes" if is_yes else "No",
                                   note if is_yes else "—"]):
            rp = row[ci].paragraphs[0]
            rr = rp.add_run(val)
            rr.bold      = ci == 1 and is_yes
            rr.font.name = _FONT_MAIN
            rr.font.size = Pt(10)
            if ci == 1 and is_yes:
                rr.font.color.rgb = _COPPER

    _set_all_col_widths(tbl, [2.0, 1.0, 3.5])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _add_parameter_review_section(doc: Document, parameters: list[dict[str, Any]]) -> None:
    doc.add_heading("4. Parameter Review and Proposed Changes", level=1)
    if not parameters:
        doc.add_paragraph("No parameters are present in this submission.")
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    _set_table_outer_border(table, _COPPER_HEX, sz=8)
    _set_table_inner_borders(table, _LGRAY_HEX, sz=4)

    headers = [
        "Parameter",
        "Type",
        "Existing Requirement",
        "Proposed Requirement",
        "Change Status",
        "Justification",
    ]
    widths = [1.8, 0.8, 1.2, 1.2, 0.8, 1.7]

    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        _shade_cell(header_cells[idx], _COPPER_HEX, font_white=True)
        paragraph = header_cells[idx].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        run.bold = True
        run.font.name = _FONT_MAIN
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _set_cell_widths(header_cells, widths)

    for parameter in parameters:
        existing = str(parameter.get("existing_value", "") or "").strip() or "—"
        proposed = str(parameter.get("proposed_value", "") or "").strip()
        justification = str(parameter.get("justification", "") or parameter.get("remarks", "") or "").strip() or "—"
        status = "Revised" if proposed and proposed != existing else "Retained"
        proposed_display = proposed or existing
        row = table.add_row().cells
        values = [
            str(parameter.get("parameter_name", "") or "Untitled Parameter"),
            str(parameter.get("parameter_type", "") or "—"),
            existing,
            proposed_display,
            status,
            justification,
        ]
        for idx, value in enumerate(values):
            paragraph = row[idx].paragraphs[0]
            run = paragraph.add_run(value)
            run.font.name = _FONT_MAIN
            run.font.size = Pt(9)
            if idx == 3 and status == "Revised":
                run.bold = True
                run.font.color.rgb = _COPPER
            if idx == 4:
                run.bold = True
                run.font.color.rgb = _COPPER if status == "Revised" else _DARK_GRAY
        _set_cell_widths(row, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _add_impact_section(
    doc: Document,
    impact_analysis: dict[str, Any] | None,
    section_map: dict[str, str],
) -> None:
    doc.add_heading("7. Impact Assessment", level=1)

    if impact_analysis:
        checklist_state_raw = impact_analysis.get("checklist_state") or impact_analysis.get("checklist_state_json")
        if checklist_state_raw:
            checklist_state = deserialize_impact_checklist_state(
                checklist_state_raw,
                impact_analysis.get("chemical_name", ""),
            )
            summary = summarize_impact_checklist_state(checklist_state)
            doc.add_paragraph(f"Chemical Name: {summary['chemical_name'] or 'Not specified'}")
            doc.add_paragraph(f"Classification: {summary['classification']}")
            doc.add_paragraph(f"Rule: {summary['rule']}")
            doc.add_paragraph(
                f"Red flags answered YES: {summary['red_yes_count']}/3 | Amber flags answered YES: {summary['amber_yes_count']}/3"
            )
            for flag in summary["flags"]:
                doc.add_paragraph(
                    f"Flag {flag['order']} ({flag['type']} / {flag['dimension']}) — {flag['answer_label']}: {flag['question']}"
                )
                doc.add_paragraph(f"Source: {flag['source']}")
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            return

        op_score = int(impact_analysis.get("operational_impact_score", 0) or 0)
        se_score = int(impact_analysis.get("safety_environment_score", 0) or 0)
        sr_score = int(impact_analysis.get("supply_risk_score", 0) or 0)
        no_substitute = int(impact_analysis.get("no_substitute_flag", 0) or 0)
        total_score = float(impact_analysis.get("impact_score_total", 0.0) or 0.0)
        final_grade = str(impact_analysis.get("impact_grade", "—") or "—")
        computed_grade = get_impact_grade(total_score, 0)
        override_applied = bool(
            no_substitute == 1
            and computed_grade in {"LOW", "MODERATE"}
            and final_grade.upper() == "HIGH"
        )

        lines = [
            f"Operational Impact Score: {op_score}",
            f"Operational Impact Note: {(impact_analysis.get('operational_note', '') or '—').strip() or '—'}",
            "",
            f"Safety / Environmental Score: {se_score}",
            f"Safety / Environmental Note: {(impact_analysis.get('safety_environment_note', '') or '—').strip() or '—'}",
            "",
            f"Supply Risk Score: {sr_score}",
            f"Supply Risk Note: {(impact_analysis.get('supply_risk_note', '') or '—').strip() or '—'}",
            "",
            f"No Substitute Available: {'Yes' if no_substitute == 1 else 'No'}",
            "",
            f"Total Impact Score: {total_score:.2f}",
            f"Final Impact Grade: {final_grade}",
        ]
        for line in lines:
            doc.add_paragraph(line)
        if override_applied:
            doc.add_paragraph("Note: No Substitute Available override applied.")
    else:
        # Backward-compatible fallback for historical drafts with text-only impact.
        body = section_map.get("Impact::operational", "").strip()
        doc.add_paragraph(body or "Impact analysis not provided.")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _add_recommendation_section(doc: Document, section_map: dict[str, str]) -> None:
    doc.add_heading("8. Approval Basis and Version Change Reason", level=1)
    rec  = section_map.get(REC_MAIN_KEY, "").strip()  or "—"
    rmks = section_map.get(REC_REMARKS_KEY, "").strip() or "—"

    tbl = doc.add_table(rows=2, cols=2)
    tbl.style = "Table Grid"
    _set_table_outer_border(tbl, _COPPER_HEX, sz=8)
    _set_table_inner_borders(tbl, _LGRAY_HEX, sz=4)

    for ri, (lbl, val) in enumerate([
        ("Approval Basis / Version Change Reason", rec),
        ("Approval Notes / Additional Remarks",  rmks),
    ]):
        cells = tbl.rows[ri].cells
        _shade_cell(cells[0], _LGRAY_HEX)
        lp = cells[0].paragraphs[0]
        lr = lp.add_run(lbl)
        lr.bold      = True
        lr.font.name = _FONT_MAIN
        lr.font.size = Pt(11)
        vp = cells[1].paragraphs[0]
        vr = vp.add_run(val)
        vr.font.name = _FONT_MAIN
        vr.font.size = Pt(11)

    for row in tbl.rows:
        _set_cell_widths(row.cells, [2.0, 4.5])

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ===========================================================================
# CHANGE LOG PAGE — final page of the master export
# ===========================================================================

def _build_change_log_page(
    doc: Document,
    changelog: list[dict[str, Any]],
    logo_path: str | Path | None = None,
) -> None:
    """Build a single-page Change Log summarising all version increments.

    Each row in changelog is expected to have:
        spec_number, chemical_name, spec_version, version_label,
        action_time, user_name, remarks
    """
    logo_path = _find_logo(logo_path)

    # ── Header ───────────────────────────────────────────────────────────────
    hdr_tbl = doc.add_table(rows=1, cols=2)
    _set_table_no_borders(hdr_tbl)

    logo_cell  = hdr_tbl.rows[0].cells[0]
    title_cell = hdr_tbl.rows[0].cells[1]

    if logo_path and logo_path.exists():
        lp = logo_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = lp.add_run()
        r.add_picture(str(logo_path), width=Inches(0.55))
    else:
        logo_cell.paragraphs[0].text = "ONGC"

    tp = title_cell.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tp.paragraph_format.space_before = Pt(4)
    r = tp.add_run("CORPORATE SPECIFICATIONS — CHANGE LOG")
    r.bold = True
    r.font.name = _FONT_TITLE
    r.font.size = Pt(14)
    r.font.color.rgb = _BLACK

    _set_cell_widths(hdr_tbl.rows[0].cells, [0.65, 5.85])

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)

    # Sub-header line
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(6)
    sr = sub.add_run(
        f"Generated: {datetime.now().strftime('%d %b %Y')}  ·  "
        f"{len(changelog)} change event(s) recorded"
    )
    sr.font.name  = _FONT_MAIN
    sr.font.size  = Pt(10)
    sr.italic     = True
    sr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    if not changelog:
        p = doc.add_paragraph("No version changes recorded yet.")
        p.paragraph_format.space_before = Pt(6)
        return

    # ── Change log table ─────────────────────────────────────────────────────
    # Columns: S.No. | Spec No. | Chemical Name | Version | Change | Date | By
    col_headers = ["S. No.", "Spec No.", "Chemical Name", "Ver.", "Change / Remarks", "Date", "By"]
    col_widths  = [0.35,      1.10,       1.80,            0.35,   2.10,               0.65,   0.15]
    # total ≈ 6.50"

    tbl = doc.add_table(rows=1, cols=len(col_headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_format_a_borders(tbl)

    # Header row
    hdr_cells = tbl.rows[0].cells
    for ci, hdr_text in enumerate(col_headers):
        _shade_cell(hdr_cells[ci], "1F2937")   # dark header
        hp = hdr_cells[ci].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(hdr_text)
        hr.bold       = True
        hr.font.name  = _FONT_MAIN
        hr.font.size  = Pt(9)
        hr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _set_cell_widths(hdr_cells, col_widths)

    # Data rows — most recent events first (already ordered DESC by repo query)
    for idx, entry in enumerate(changelog, start=1):
        row_cells = tbl.add_row().cells

        # Friendly action label
        action = entry.get("action", "") or ""
        if "VERSION_INCREMENT" in action:
            change_desc = entry.get("remarks") or "Parameter data updated"
        elif "PUBLISHED" in action:
            change_desc = "Published to master data"
        elif "PDF_IMPORT" in action or "CREATE_DRAFT" in action:
            change_desc = "Initial ingest / creation"
        else:
            change_desc = action.replace("_", " ").title()

        # Format date
        raw_dt = entry.get("action_time", "")
        date_str = raw_dt[:10] if raw_dt else "—"

        ver_str = entry.get("version_label") or f"v{entry.get('spec_version', '?')}"
        row_data = [
            str(idx),
            entry.get("spec_number", "—"),
            entry.get("chemical_name", "—"),
            ver_str,
            change_desc,
            date_str,
            entry.get("user_name", "—")[:12],
        ]

        for ci, val in enumerate(row_data):
            rp = row_cells[ci].paragraphs[0]
            if ci in (0, 3):
                rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = rp.add_run(val)
            rr.font.name  = _FONT_MAIN
            rr.font.size  = Pt(9)
            rr.font.color.rgb = _BLACK

        _set_cell_widths(row_cells, col_widths)

    sp2 = doc.add_paragraph()
    sp2.paragraph_format.space_after = Pt(6)


# ---------------------------------------------------------------------------
# Footer
# ---------------------------------------------------------------------------

def _add_footer(doc: Document, draft: dict[str, Any]) -> None:
    section = doc.sections[0]
    _set_footer_page_number(section.footer)


def _add_master_footer(doc: Document, spec_count: int) -> None:
    """Footer for the multi-spec master export document."""
    section = doc.sections[0]
    _set_footer_page_number(section.footer)


def _set_footer_page_number(footer) -> None:
    """Set a right-aligned footer PAGE field similar to the cleaned reference file."""
    for para in list(footer.paragraphs):
        para._element.getparent().remove(para._element)
    for tbl in list(footer.tables):
        tbl._element.getparent().remove(tbl._element)

    para = footer.add_paragraph()
    p = para._p
    pPr = p.get_or_add_pPr()

    p_style = OxmlElement("w:pStyle")
    p_style.set(qn("w:val"), "Footer")
    pPr.append(p_style)

    frame = OxmlElement("w:framePr")
    frame.set(qn("w:wrap"), "none")
    frame.set(qn("w:vAnchor"), "text")
    frame.set(qn("w:hAnchor"), "margin")
    frame.set(qn("w:xAlign"), "right")
    frame.set(qn("w:y"), "1")
    pPr.append(frame)

    def _append_fld_char(kind: str) -> None:
        run = para.add_run()
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), kind)
        run._r.append(fld)

    _append_fld_char("begin")
    instr_run = para.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    instr_run._r.append(instr)
    _append_fld_char("separate")
    _append_fld_char("end")

    pad = footer.add_paragraph()
    pad.style = "Footer"
    pad.paragraph_format.right_indent = Twips(360)


# ===========================================================================
# Page / Style helpers
# ===========================================================================

def _configure_page(doc: Document) -> None:
    sec = doc.sections[0]
    # Match cleaned document page geometry exactly (twips).
    sec.page_width      = Twips(11_909)
    sec.page_height     = Twips(16_834)
    sec.left_margin     = Twips(1_296)
    sec.right_margin    = Twips(850)
    sec.top_margin      = Twips(1_116)
    sec.bottom_margin   = Twips(1_224)
    sec.header_distance = Twips(607)
    sec.footer_distance = Twips(720)


def _apply_base_styles(doc: Document) -> None:
    """Apply Times New Roman 11pt to Normal style, configure headings."""
    normal = doc.styles["Normal"]
    _apply_style_font(normal, _FONT_MAIN, 11, False, _BLACK)

    for level, size, bold, color in [
        ("Heading 1", 12, True,  _COPPER),
        ("Heading 2", 11, True,  _DARK_GRAY),
        ("Heading 3", 11, False, _DARK_GRAY),
    ]:
        try:
            h = doc.styles[level]
            _apply_style_font(h, _FONT_MAIN, size, bold, color)
        except KeyError:
            pass


def _apply_style_font(style, font_name: str, size_pt: int, bold: bool, color: RGBColor) -> None:
    """Force style rFonts to concrete family values (not theme references)."""
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.color.rgb = color

    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)

    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), font_name)
    for attr in ("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme"):
        key = qn(f"w:{attr}")
        if key in rfonts.attrib:
            del rfonts.attrib[key]

    for tag in ("sz", "szCs"):
        size_el = rpr.find(qn(f"w:{tag}"))
        if size_el is None:
            size_el = OxmlElement(f"w:{tag}")
            rpr.append(size_el)
        size_el.set(qn("w:val"), str(size_pt * 2))


# ===========================================================================
# Low-level XML / border helpers
# ===========================================================================

def _shade_cell(cell, fill_hex: str, font_white: bool = False) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for ex in tcPr.findall(qn("w:shd")):
        tcPr.remove(ex)
    tcPr.append(parse_xml(
        f'<w:shd {nsdecls("w")} '
        f'w:val="clear" w:color="auto" w:fill="{fill_hex}"/>'
    ))
    if font_white:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _border_xml(color: str, sz: int, val: str = "single", space: int = 0) -> str:
    return (
        f'w:val="{val}" w:sz="{sz}" '
        f'w:space="{space}" w:color="{color}"'
    )


def _read_border_attrs(tbl_borders, side: str, fallback_color: str, fallback_sz: int) -> tuple[str, int, str, int]:
    """Read one side from existing tblBorders, with fallback defaults."""
    side_el = tbl_borders.find(qn(f"w:{side}")) if tbl_borders is not None else None
    if side_el is None:
        return "single", fallback_sz, fallback_color, 0
    val = side_el.get(qn("w:val"), "single")
    sz = int(side_el.get(qn("w:sz"), str(fallback_sz)))
    color = side_el.get(qn("w:color"), fallback_color)
    space = int(side_el.get(qn("w:space"), "0"))
    return val, sz, color, space


def _set_table_outer_border(table, color: str, sz: int = 6) -> None:
    """Apply a uniform outer+inner border to the table via tblPr."""
    tbl   = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.add_tblPr()
    for ex in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(ex)
    bdr = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top    {_border_xml(color, sz)}/>'
        f'  <w:left   {_border_xml(color, sz)}/>'
        f'  <w:bottom {_border_xml(color, sz)}/>'
        f'  <w:right  {_border_xml(color, sz)}/>'
        f'  <w:insideH {_border_xml(color, sz)}/>'
        f'  <w:insideV {_border_xml(color, sz)}/>'
        f'</w:tblBorders>'
    )
    tblPr.append(bdr)


def _set_table_inner_borders(table, color: str, sz: int = 4) -> None:
    """Set only the inner borders of a table while preserving outer borders."""
    tbl   = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.add_tblPr()
    existing = tblPr.find(qn("w:tblBorders"))
    top_val, top_sz, top_color, top_space = _read_border_attrs(existing, "top", color, sz)
    left_val, left_sz, left_color, left_space = _read_border_attrs(existing, "left", color, sz)
    bottom_val, bottom_sz, bottom_color, bottom_space = _read_border_attrs(existing, "bottom", color, sz)
    right_val, right_sz, right_color, right_space = _read_border_attrs(existing, "right", color, sz)

    for ex in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(ex)
    bdr = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top    {_border_xml(top_color, top_sz, top_val, top_space)}/>'
        f'  <w:left   {_border_xml(left_color, left_sz, left_val, left_space)}/>'
        f'  <w:bottom {_border_xml(bottom_color, bottom_sz, bottom_val, bottom_space)}/>'
        f'  <w:right  {_border_xml(right_color, right_sz, right_val, right_space)}/>'
        f'  <w:insideH {_border_xml(color, sz)}/>'
        f'  <w:insideV {_border_xml(color, sz)}/>'
        f'</w:tblBorders>'
    )
    tblPr.append(bdr)


def _set_table_format_a_borders(table) -> None:
    """Apply Format A cell borders: single, sz=4 (0.5pt), black — matching source doc."""
    _set_table_outer_border(table, "000000", sz=4)


def _set_table_no_borders(table) -> None:
    """Remove table border definitions (for layout tables)."""
    tbl   = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.add_tblPr()
    for ex in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(ex)
    for ex in tblPr.findall(qn("w:jc")):
        tblPr.remove(ex)


def _set_cell_widths(cells, widths_inches: list[float]) -> None:
    for idx, cell in enumerate(cells):
        if idx < len(widths_inches):
            cell.width = Inches(widths_inches[idx])


def _set_cell_widths_dxa(cells, widths_dxa: list[int]) -> None:
    for idx, cell in enumerate(cells):
        if idx >= len(widths_dxa):
            continue
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcw = tcPr.find(qn("w:tcW"))
        if tcw is None:
            tcw = OxmlElement("w:tcW")
            tcPr.append(tcw)
        tcw.set(qn("w:w"), str(int(widths_dxa[idx])))
        tcw.set(qn("w:type"), "dxa")


def _set_table_grid_dxa(table, widths_dxa: list[int]) -> None:
    tbl = table._tbl
    tbl_grid = tbl.tblGrid
    if tbl_grid is not None:
        tbl.remove(tbl_grid)
    new_grid = OxmlElement("w:tblGrid")
    for w in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(w)))
        new_grid.append(col)
    # Keep ordering: after tblPr if present, before first row.
    if tbl.tblPr is not None:
        tbl.insert(1, new_grid)
    else:
        tbl.insert(0, new_grid)


def _set_table_width_dxa(table, width_dxa: int) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.add_tblPr()
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(int(width_dxa)))
    tblW.set(qn("w:type"), "dxa" if width_dxa > 0 else "auto")


def _set_table_indent_dxa(table, indent_dxa: int) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.add_tblPr()
    for ex in tblPr.findall(qn("w:jc")):
        tblPr.remove(ex)
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(int(indent_dxa)))
    tblInd.set(qn("w:type"), "dxa")


def _set_all_col_widths(table, widths_inches: list[float]) -> None:
    for row in table.rows:
        _set_cell_widths(row.cells, widths_inches)
