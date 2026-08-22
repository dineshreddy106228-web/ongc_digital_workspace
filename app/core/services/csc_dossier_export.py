"""Enterprise Word dossier export for one corporate specification.

The CSC catalogue can map one specification record to several register chemicals.
This exporter therefore treats the dossier as a controlled technical document: it
shows every covered chemical, the live requirements, supporting technical data and
the version trail without reproducing empty legacy-workflow comparison columns.
"""

from __future__ import annotations

import io
import math
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.core.services.csc_export import _find_logo


# compact_reference_guide preset, resolved to concrete document tokens.
PAGE_WIDTH_IN = 8.5
PAGE_HEIGHT_IN = 11.0
MARGIN_IN = 1.0
HEADER_FOOTER_IN = 0.492
CONTENT_WIDTH_DXA = 9_360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_DXA = 80
CELL_MARGIN_BOTTOM_DXA = 80
CELL_MARGIN_START_DXA = 120
CELL_MARGIN_END_DXA = 120

FONT_BODY = "Calibri"
COLOR_INK = "0B2545"
COLOR_HEADING = "2E74B5"
COLOR_HEADING_DARK = "1F4D78"
COLOR_ACCENT = "0E766E"  # named ONGC enterprise accent override
COLOR_ACCENT_DARK = "075A54"
COLOR_GOLD = "9A6B16"
COLOR_TEXT = "1F2937"
COLOR_MUTED = "667085"
COLOR_BORDER = "C9D3DF"
COLOR_HEADER_FILL = "E8EEF5"
COLOR_SOFT_FILL = "F4F6F9"
COLOR_TEAL_FILL = "E6F3F1"
COLOR_WHITE = "FFFFFF"
COLOR_RISK = "9B1C1C"

EM_DASH = "—"


def build_enterprise_dossier(
    context: dict[str, Any],
    logo_path: str | Path | None = None,
) -> bytes:
    """Build a polished, read-only corporate specification dossier."""

    document = Document()
    _configure_document(document)
    _configure_styles(document)
    _configure_properties(document, context)
    _enable_update_fields(document)

    _build_cover(document, context, logo_path)
    _start_controlled_page(document)
    _build_executive_summary(document, context)
    _start_controlled_page(document)
    _build_parameter_register(document, context)
    _start_controlled_page(document)
    _build_narrative_sections(document, context)
    _build_technical_profile(document, context)
    _build_impact_and_issues(document, context)
    _build_governance(document, context)
    _configure_header_footer(document, context, logo_path)

    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(PAGE_WIDTH_IN)
    section.page_height = Inches(PAGE_HEIGHT_IN)
    section.top_margin = Inches(MARGIN_IN)
    section.right_margin = Inches(MARGIN_IN)
    section.bottom_margin = Inches(MARGIN_IN)
    section.left_margin = Inches(MARGIN_IN)
    section.header_distance = Inches(HEADER_FOOTER_IN)
    section.footer_distance = Inches(HEADER_FOOTER_IN)


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    _set_style_font(normal, FONT_BODY, 11, COLOR_TEXT, bold=False)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, COLOR_HEADING, 18, 10),
        "Heading 2": (13, COLOR_HEADING, 14, 7),
        "Heading 3": (12, COLOR_HEADING_DARK, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[name]
        _set_style_font(style, FONT_BODY, size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("Header", "Footer"):
        style = document.styles[name]
        _set_style_font(style, FONT_BODY, 8.5, COLOR_MUTED, bold=False)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)


def _configure_properties(document: Document, context: dict[str, Any]) -> None:
    draft = context.get("draft") or {}
    document.core_properties.title = (
        f"{_text(draft.get('spec_number'))} — {_text(draft.get('chemical_name'))}"
    )
    document.core_properties.subject = "ONGC Corporate Specification Dossier"
    document.core_properties.author = "ONGC Corporate Chemistry"
    document.core_properties.category = "Controlled Technical Document"
    document.core_properties.keywords = "ONGC, corporate specification, oil field chemicals, CSC"
    document.core_properties.comments = (
        "System-generated controlled dossier from ONGC Digital Workspace."
    )


def _configure_header_footer(
    document: Document,
    context: dict[str, Any],
    logo_path: str | Path | None,
) -> None:
    draft = context.get("draft") or {}
    spec_number = _text(draft.get("spec_number"))
    for section in document.sections:
        section.different_first_page_header_footer = True
        for header in (
            section.first_page_header,
            section.header,
            section.even_page_header,
        ):
            header.is_linked_to_previous = False
            _populate_running_header(header, spec_number, logo_path)
        for footer in (
            section.first_page_footer,
            section.footer,
            section.even_page_footer,
        ):
            footer.is_linked_to_previous = False
            _populate_running_footer(footer, spec_number)


def _start_controlled_page(document: Document) -> None:
    """Begin a new-page section so running matter is stable across office suites."""
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = Inches(PAGE_WIDTH_IN)
    section.page_height = Inches(PAGE_HEIGHT_IN)
    section.top_margin = Inches(MARGIN_IN)
    section.right_margin = Inches(MARGIN_IN)
    section.bottom_margin = Inches(MARGIN_IN)
    section.left_margin = Inches(MARGIN_IN)
    section.header_distance = Inches(HEADER_FOOTER_IN)
    section.footer_distance = Inches(HEADER_FOOTER_IN)


def _populate_running_header(header, spec_number: str, logo_path: str | Path | None) -> None:
    _clear_container(header)
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    logo = _find_logo(logo_path)
    left, right = header_table.rows[0].cells
    if logo and logo.exists():
        run = left.paragraphs[0].add_run()
        run.add_picture(str(logo), width=Inches(0.42))
        _set_last_picture_alt(run, "ONGC Corporate Chemistry emblem")
    else:
        _add_run(left.paragraphs[0], "ONGC", bold=True, color=COLOR_ACCENT_DARK, size=10)
    right_p = right.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(right_p, "CORPORATE SPECIFICATION DOSSIER", bold=True, color=COLOR_INK, size=8.5)
    _add_run(right_p, f"\n{spec_number}", color=COLOR_MUTED, size=8)
    _finish_table(header_table, [1_050, 8_310], indent_dxa=0, borders=False)


def _populate_running_footer(footer, spec_number: str) -> None:
    _clear_container(footer)
    footer_table = footer.add_table(rows=1, cols=3, width=Inches(6.5))
    cells = footer_table.rows[0].cells
    _add_run(cells[0].paragraphs[0], spec_number, bold=True, color=COLOR_MUTED, size=8)
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(cells[1].paragraphs[0], "CONTROLLED COPY", bold=True, color=COLOR_ACCENT_DARK, size=8)
    cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(cells[2].paragraphs[0], "Page ", color=COLOR_MUTED, size=8)
    _append_field(cells[2].paragraphs[0], "PAGE", "1")
    _add_run(cells[2].paragraphs[0], " of ", color=COLOR_MUTED, size=8)
    _append_field(cells[2].paragraphs[0], "NUMPAGES", "1")
    _finish_table(footer_table, [3_120, 3_120, 3_120], indent_dxa=0, borders=False)


def _build_cover(
    document: Document,
    context: dict[str, Any],
    logo_path: str | Path | None,
) -> None:
    draft = context.get("draft") or {}
    generated_at = datetime.now(timezone.utc).astimezone()
    document.add_paragraph().paragraph_format.space_after = Pt(34)

    logo = _find_logo(logo_path)
    logo_p = document.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo and logo.exists():
        run = logo_p.add_run()
        run.add_picture(str(logo), width=Inches(0.92))
        _set_last_picture_alt(run, "ONGC Corporate Chemistry emblem")
    else:
        _add_run(logo_p, "ONGC", bold=True, color=COLOR_ACCENT_DARK, size=18)
    logo_p.paragraph_format.space_after = Pt(16)

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(8)
    _add_run(kicker, "CORPORATE SPECIFICATION DOSSIER", bold=True, color=COLOR_ACCENT, size=11)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(7)
    _add_run(title, _text(draft.get("chemical_name")), bold=True, color=COLOR_INK, size=25)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    _add_run(
        subtitle,
        f"{_text(draft.get('spec_number'))}  |  {_text(draft.get('version_display'))}",
        bold=True,
        color=COLOR_HEADING_DARK,
        size=13,
    )

    rule = document.add_paragraph()
    rule.paragraph_format.space_after = Pt(18)
    _set_paragraph_bottom_border(rule, COLOR_ACCENT, size=18)

    summary = _summary_map(context)
    cover_rows = [
        ("Document class", "Controlled Technical Document"),
        ("Category", summary.get("Category", EM_DASH)),
        ("Specification status", "Current / Published"),
        ("Corporate register", summary.get("On Corporate Register", EM_DASH)),
        ("Prepared by", summary.get("Prepared By", EM_DASH)),
        ("Reviewed by", summary.get("Reviewed By", EM_DASH)),
        ("Last revised", summary.get("Last Revised", EM_DASH)),
        ("Generated", generated_at.strftime("%d %b %Y, %H:%M %Z")),
    ]
    _add_label_value_table(document, cover_rows, label_width=2_700)

    document.add_paragraph().paragraph_format.space_after = Pt(10)
    _add_callout(
        document,
        "Purpose",
        "This dossier consolidates the current corporate specification, covered register "
        "chemicals, technical requirements, material information, impact controls and "
        "version history into one read-only review and reference document.",
    )

    controlled = document.add_paragraph()
    controlled.alignment = WD_ALIGN_PARAGRAPH.CENTER
    controlled.paragraph_format.space_before = Pt(16)
    _add_run(
        controlled,
        "Generated from the live ONGC Digital Workspace record. Uncontrolled when printed.",
        italic=True,
        color=COLOR_MUTED,
        size=8.5,
    )


def _build_executive_summary(document: Document, context: dict[str, Any]) -> None:
    draft = context.get("draft") or {}
    summary = _summary_map(context)
    parameters = context.get("parameter_rows") or []
    covered = context.get("covered_chemicals") or []

    document.add_heading("1. Executive Summary", level=1)
    intro = document.add_paragraph()
    _add_run(
        intro,
        f"{_text(draft.get('chemical_name'))} is maintained under "
        f"{_text(draft.get('spec_number'))}. This controlled release contains "
        f"{len(parameters)} specification parameters and covers {len(covered) or 1} "
        "corporate-register chemical entries.",
        size=11,
    )

    metrics = [
        ("PARAMETERS", str(len(parameters))),
        ("VERSION", _text(draft.get("version_display"))),
        ("IMPACT", summary.get("Impact Classification", EM_DASH)),
        ("REGISTER", summary.get("On Corporate Register", EM_DASH)),
    ]
    table = document.add_table(rows=2, cols=4)
    for index, (label, value) in enumerate(metrics):
        label_p = table.rows[0].cells[index].paragraphs[0]
        label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(label_p, label, bold=True, color=COLOR_MUTED, size=8)
        value_p = table.rows[1].cells[index].paragraphs[0]
        value_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(value_p, value, bold=True, color=COLOR_INK, size=13)
        _shade_cell(table.rows[1].cells[index], COLOR_TEAL_FILL)
    _finish_table(table, [2_340, 2_340, 2_340, 2_340], borders=True)

    document.add_heading("Document Control", level=2)
    control_rows = [
        ("Specification", summary.get("Specification", _text(draft.get("spec_number")))),
        ("Chemical", summary.get("Chemical", _text(draft.get("chemical_name")))),
        ("Version", summary.get("Version", _text(draft.get("version_display")))),
        ("Test procedure", _text(draft.get("test_procedure"))),
        ("Material reference", _text(draft.get("material_code"))),
        ("Standard testing time", summary.get("Standard Testing Time", EM_DASH)),
        ("Prepared by", summary.get("Prepared By", EM_DASH)),
        ("Reviewed by", summary.get("Reviewed By", EM_DASH)),
    ]
    _add_label_value_table(document, control_rows)

    document.add_heading("Covered Register Chemicals", level=2)
    if covered:
        table = document.add_table(rows=1, cols=4)
        headers = ["S. No.", "Chemical / Grade", "Material Code", "Testing Time"]
        for index, label in enumerate(headers):
            _set_header_cell(table.rows[0].cells[index], label)
        for index, item in enumerate(covered, start=1):
            cells = table.add_row().cells
            values = [
                str(index),
                _text(item.get("chemical_name")),
                _text(item.get("material_code")),
                _testing_time(item.get("standard_days")),
            ]
            for cell_index, value in enumerate(values):
                alignment = WD_ALIGN_PARAGRAPH.CENTER if cell_index in {0, 2, 3} else WD_ALIGN_PARAGRAPH.LEFT
                _set_cell_value(cells[cell_index], value, alignment=alignment, size=9.5)
            if index % 2 == 0:
                for cell in cells:
                    _shade_cell(cell, COLOR_SOFT_FILL)
        _finish_table(table, [650, 5_000, 1_950, 1_760], header=True)
    else:
        document.add_paragraph("No covered-register mapping is available.")


def _build_parameter_register(document: Document, context: dict[str, Any]) -> None:
    rows = context.get("parameter_rows") or []
    document.add_heading("2. Specification Requirements", level=1)
    lead = document.add_paragraph()
    _add_run(
        lead,
        "The following requirements constitute the current controlled parameter set. "
        "Parameter classifications are shown below each parameter name.",
        color=COLOR_MUTED,
        size=10,
    )
    if not rows:
        _add_callout(document, "Status", "No specification parameters are recorded.")
        return

    number = 1
    for chunk_index, chunk in enumerate(_parameter_chunks(rows)):
        if chunk_index:
            _start_controlled_page(document)
        table = _add_parameter_table(document, chunk, number)
        number += len(chunk)
        _finish_table(table, [560, 2_450, 3_380, 900, 2_070], header=True)


def _add_parameter_table(document: Document, rows: list[dict[str, Any]], start_number: int):
    table = document.add_table(rows=1, cols=5)
    headers = ["S. No.", "Parameter", "Required Value", "Unit", "Method / Conditions"]
    for index, label in enumerate(headers):
        _set_header_cell(table.rows[0].cells[index], label, dark=True)

    for offset, parameter in enumerate(rows):
        number = start_number + offset
        cells = table.add_row().cells
        _set_cell_value(cells[0], str(number), alignment=WD_ALIGN_PARAGRAPH.CENTER, size=9)

        parameter_p = cells[1].paragraphs[0]
        _add_run(parameter_p, _text(parameter.get("parameter_name")), bold=True, size=9.5)
        parameter_type = _meaningful(parameter.get("parameter_type"))
        if parameter_type:
            _add_run(parameter_p, f"\n{parameter_type}", italic=True, color=COLOR_ACCENT_DARK, size=8.5)

        _set_cell_value(cells[2], _text(parameter.get("final_requirement")), size=9.5)
        _set_cell_value(
            cells[3],
            _text(parameter.get("unit_of_measure")),
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            size=9,
        )
        technical = []
        conditions = _meaningful(parameter.get("conditions"))
        procedure = _meaningful(parameter.get("procedure_text"))
        if conditions:
            technical.append(f"Conditions: {conditions}")
        if procedure:
            technical.append(f"Method: {procedure}")
        _set_cell_value(cells[4], "\n".join(technical) or EM_DASH, size=8.5)

        if number % 2 == 0:
            for cell in cells:
                _shade_cell(cell, COLOR_SOFT_FILL)
    return table


def _parameter_chunks(rows: list[dict[str, Any]]) -> list[list[dict[str, Any]]]:
    """Estimate page-safe table segments to avoid renderer-dependent overflow."""
    chunks: list[list[dict[str, Any]]] = []
    current: list[dict[str, Any]] = []
    current_weight = 0

    for parameter in rows:
        budget = 23 if not chunks else 27
        weight = _parameter_row_weight(parameter)
        if current and current_weight + weight > budget:
            chunks.append(current)
            current = []
            current_weight = 0
        current.append(parameter)
        current_weight += weight
    if current:
        chunks.append(current)

    # Avoid a visually weak final page containing only one or two short rows.
    if len(chunks) > 1 and len(chunks[-1]) < 3 and len(chunks[-2]) > 4:
        while len(chunks[-1]) < 3 and len(chunks[-2]) > 4:
            chunks[-1].insert(0, chunks[-2].pop())
    return chunks


def _parameter_row_weight(parameter: dict[str, Any]) -> int:
    name = _text(parameter.get("parameter_name"))
    requirement = _text(parameter.get("final_requirement"))
    procedure = _text(parameter.get("procedure_text"))
    conditions = _text(parameter.get("conditions"))
    return max(
        2,
        math.ceil(len(name) / 34) + 1,
        math.ceil(len(requirement) / 48),
        math.ceil((len(procedure) + len(conditions)) / 28),
    )


def _build_narrative_sections(document: Document, context: dict[str, Any]) -> None:
    document.add_heading("3. Specification Narrative", level=1)
    rows = [row for row in (context.get("section_rows") or []) if _meaningful(row.get("text"))]
    if not rows:
        document.add_paragraph("No narrative sections are recorded for this specification.")
        return

    for row in rows:
        document.add_heading(_text(row.get("label")), level=2)
        current = _meaningful(row.get("text")) or EM_DASH
        previous = _meaningful(row.get("source_text"))
        if previous and previous != current:
            table = document.add_table(rows=1, cols=2)
            _set_header_cell(table.rows[0].cells[0], "Previous Release")
            _set_header_cell(table.rows[0].cells[1], "Current Release")
            cells = table.add_row().cells
            _set_cell_value(cells[0], previous)
            _set_cell_value(cells[1], current)
            _finish_table(table, [4_680, 4_680], header=True)
        else:
            document.add_paragraph(current)


def _build_technical_profile(document: Document, context: dict[str, Any]) -> None:
    document.add_heading("4. Material and Handling Profile", level=1)
    groups = [
        ("Material Identity", context.get("master_rows") or []),
        ("Material Properties", context.get("material_property_rows") or []),
        ("Storage and Handling", context.get("storage_rows") or []),
    ]
    any_rows = False
    for title, rows in groups:
        rows = [row for row in rows if _meaningful(row.get("value"))]
        if not rows:
            continue
        any_rows = True
        document.add_heading(title, level=2)
        _add_context_rows(document, rows)
    if not any_rows:
        document.add_paragraph("No additional material or handling data is recorded.")


def _build_impact_and_issues(document: Document, context: dict[str, Any]) -> None:
    document.add_heading("5. Impact and Issue Controls", level=1)
    impact_rows = [row for row in (context.get("impact_rows") or []) if _meaningful(row.get("value"))]
    summary_labels = {
        "Impact Classification", "Decision Rule", "Red YES Count", "Amber YES Count", "Flags Answered"
    }
    summary = [row for row in impact_rows if row.get("label") in summary_labels]
    flags = [row for row in impact_rows if row.get("label") not in summary_labels]

    if summary:
        document.add_heading("Impact Summary", level=2)
        _add_context_rows(document, summary)
    if flags:
        if len(flags) >= 6:
            _start_controlled_page(document)
        document.add_heading("Assessment Flags", level=2)
        _add_assessment_flags_table(document, flags)

    issues = context.get("issue_rows") or []
    document.add_heading("Issue Register", level=2)
    table = document.add_table(rows=1, cols=3)
    for index, label in enumerate(["Issue Domain", "Status", "Recorded Note"]):
        _set_header_cell(table.rows[0].cells[index], label)
    for index, issue in enumerate(issues, start=1):
        cells = table.add_row().cells
        present = bool(issue.get("is_present"))
        values = [
            _text(issue.get("label")),
            "OPEN" if present else "NO ISSUE RECORDED",
            _text(issue.get("note")) if present else EM_DASH,
        ]
        _set_cell_value(cells[0], values[0], size=9.5)
        _set_cell_value(cells[1], values[1], alignment=WD_ALIGN_PARAGRAPH.CENTER, size=8.5)
        _set_cell_value(cells[2], values[2], size=9.5)
        status_run = cells[1].paragraphs[0].runs[0]
        status_run.bold = True
        status_run.font.color.rgb = _rgb(COLOR_RISK if present else COLOR_ACCENT_DARK)
        if index % 2 == 0:
            for cell in cells:
                _shade_cell(cell, COLOR_SOFT_FILL)
    _finish_table(table, [2_150, 2_050, 5_160], header=True)


def _build_governance(document: Document, context: dict[str, Any]) -> None:
    document.add_heading("6. Governance and Revision Record", level=1)
    versions = context.get("versions") or []
    if versions:
        table = document.add_table(rows=1, cols=5)
        for index, label in enumerate(["Version", "Date", "Action", "Recorded By", "Reason / Note"]):
            _set_header_cell(table.rows[0].cells[index], label, dark=True)
        for index, version in enumerate(versions, start=1):
            cells = table.add_row().cells
            created_at = version.get("created_at")
            date_text = created_at.strftime("%d %b %Y") if hasattr(created_at, "strftime") else _text(created_at)
            values = [
                f"v{_text(version.get('version'))}",
                date_text,
                _humanize_action(version.get("action")),
                _text(version.get("created_by")),
                _text(version.get("reason")),
            ]
            for cell_index, value in enumerate(values):
                alignment = WD_ALIGN_PARAGRAPH.CENTER if cell_index in {0, 1} else WD_ALIGN_PARAGRAPH.LEFT
                _set_cell_value(cells[cell_index], value, alignment=alignment, size=8.7)
            if version.get("is_current"):
                _shade_cell(cells[0], COLOR_TEAL_FILL)
                cells[0].paragraphs[0].runs[0].font.color.rgb = _rgb(COLOR_ACCENT_DARK)
                cells[0].paragraphs[0].runs[0].bold = True
            if index % 2 == 0:
                for cell in cells[1:]:
                    _shade_cell(cell, COLOR_SOFT_FILL)
        _finish_table(table, [1_000, 1_350, 1_650, 1_850, 3_510], header=True)
    else:
        document.add_paragraph("No version-history entries are recorded.")

    notes = _meaningful(context.get("latest_review_notes"))
    document.add_heading("Review Note", level=2)
    document.add_paragraph(notes or "No prior administrative review note is recorded.")


def _add_context_rows(document: Document, rows: list[dict[str, Any]]) -> None:
    has_baseline = any(_meaningful(row.get("source_value")) for row in rows)
    if not has_baseline:
        _add_label_value_table(
            document,
            [(_text(row.get("label")), _text(row.get("value"))) for row in rows],
        )
        return

    table = document.add_table(rows=1, cols=4)
    for index, label in enumerate(["Field", "Previous Release", "Current Release", "Status"]):
        _set_header_cell(table.rows[0].cells[index], label)
    for index, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        values = [
            _text(row.get("label")),
            _text(row.get("source_value")),
            _text(row.get("value")),
            _text(row.get("change_status")),
        ]
        for cell_index, value in enumerate(values):
            alignment = WD_ALIGN_PARAGRAPH.CENTER if cell_index == 3 else WD_ALIGN_PARAGRAPH.LEFT
            _set_cell_value(cells[cell_index], value, alignment=alignment, size=9)
        if index % 2 == 0:
            for cell in cells:
                _shade_cell(cell, COLOR_SOFT_FILL)
    _finish_table(table, [2_050, 2_550, 2_900, 1_860], header=True)


def _add_assessment_flags_table(document: Document, rows: list[dict[str, Any]]) -> None:
    table = document.add_table(rows=1, cols=2)
    _set_header_cell(table.rows[0].cells[0], "Assessment Dimension")
    _set_header_cell(table.rows[0].cells[1], "Recorded Response")
    for index, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        _set_cell_value(cells[0], _text(row.get("label")), bold=True, size=9.5)
        _set_cell_value(cells[1], _text(row.get("value")), size=9.5)
        if index % 2 == 0:
            for cell in cells:
                _shade_cell(cell, COLOR_SOFT_FILL)
    _finish_table(table, [3_400, 5_960], header=True)


def _add_label_value_table(
    document: Document,
    rows: Iterable[tuple[str, Any]],
    *,
    label_width: int = 2_700,
) -> None:
    rows = list(rows)
    table = document.add_table(rows=0, cols=2)
    for index, (label, value) in enumerate(rows, start=1):
        cells = table.add_row().cells
        _shade_cell(cells[0], COLOR_HEADER_FILL)
        _set_cell_value(cells[0], _text(label), bold=True, color=COLOR_INK, size=9.5)
        _set_cell_value(cells[1], _text(value), size=9.5)
        if index % 2 == 0:
            _shade_cell(cells[1], COLOR_SOFT_FILL)
    _finish_table(table, [label_width, CONTENT_WIDTH_DXA - label_width], borders=True)


def _add_callout(document: Document, label: str, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _shade_cell(cell, COLOR_TEAL_FILL)
    paragraph = cell.paragraphs[0]
    _add_run(paragraph, f"{label}: ", bold=True, color=COLOR_ACCENT_DARK, size=10)
    _add_run(paragraph, text, color=COLOR_TEXT, size=10)
    _finish_table(table, [CONTENT_WIDTH_DXA], borders=True, border_color=COLOR_ACCENT)


def _finish_table(
    table,
    widths_dxa: list[int],
    *,
    indent_dxa: int = TABLE_INDENT_DXA,
    header: bool = False,
    borders: bool = True,
    border_color: str = COLOR_BORDER,
) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Dossier table widths must total {CONTENT_WIDTH_DXA} DXA")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    for element in tbl_pr.findall(qn("w:tblW")):
        tbl_pr.remove(element)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

    for element in tbl_pr.findall(qn("w:tblInd")):
        tbl_pr.remove(element)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    for element in tbl_pr.findall(qn("w:tblLayout")):
        tbl_pr.remove(element)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    old_grid = tbl.tblGrid
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    tbl.insert(1, grid)

    _set_table_borders(table, border_color if borders else "nil")
    for row_index, row in enumerate(table.rows):
        _set_row_cant_split(row)
        if header and row_index == 0:
            _set_row_repeat(row)
        for index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_width_dxa(cell, widths_dxa[index])
            _set_cell_margins(cell)

def _set_table_borders(table, color: str) -> None:
    tbl_pr = table._tbl.tblPr
    for element in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(element)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        if color == "nil":
            node.set(qn("w:val"), "nil")
        else:
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), "5")
            node.set(qn("w:space"), "0")
            node.set(qn("w:color"), color)
        borders.append(node)
    tbl_pr.append(borders)


def _set_header_cell(cell, text: str, *, dark: bool = False) -> None:
    fill = COLOR_ACCENT if dark else COLOR_HEADER_FILL
    color = COLOR_WHITE if dark else COLOR_INK
    _shade_cell(cell, fill)
    _set_cell_value(cell, text, bold=True, color=color, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=9)


def _set_cell_value(
    cell,
    value: Any,
    *,
    bold: bool = False,
    color: str = COLOR_TEXT,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    size: float = 9.5,
) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    _add_run(paragraph, _text(value), bold=bold, color=color, size=size)


def _set_cell_width_dxa(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def _set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    values = {
        "top": CELL_MARGIN_TOP_DXA,
        "bottom": CELL_MARGIN_BOTTOM_DXA,
        "start": CELL_MARGIN_START_DXA,
        "end": CELL_MARGIN_END_DXA,
    }
    for side, value in values.items():
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_row_repeat(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def _set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for element in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(element)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _set_paragraph_bottom_border(paragraph, color: str, *, size: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def _set_style_font(style, name: str, size: float, color: str, *, bold: bool) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = _rgb(color)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), name)


def _add_run(
    paragraph,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    color: str = COLOR_TEXT,
    size: float = 11,
):
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.italic = italic
    run.font.name = FONT_BODY
    run.font.size = Pt(size)
    run.font.color.rgb = _rgb(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), FONT_BODY)
    r_fonts.set(qn("w:hAnsi"), FONT_BODY)
    return run


def _append_field(paragraph, instruction: str, fallback: str) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run._r.append(begin)

    instruction_run = paragraph.add_run()
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = f" {instruction} "
    instruction_run._r.append(instruction_text)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)
    _add_run(paragraph, fallback, color=COLOR_MUTED, size=8)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def _enable_update_fields(document: Document) -> None:
    settings = document.settings.element
    for element in settings.findall(qn("w:updateFields")):
        settings.remove(element)
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def _set_last_picture_alt(run, description: str) -> None:
    doc_properties = run._r.xpath(".//wp:docPr")
    if doc_properties:
        doc_properties[-1].set("descr", description)
        doc_properties[-1].set("title", description)


def _clear_container(container) -> None:
    for paragraph in list(container.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)
    for table in list(container.tables):
        table._element.getparent().remove(table._element)


def _summary_map(context: dict[str, Any]) -> dict[str, str]:
    return {
        _text(row.get("label")): _text(row.get("value"))
        for row in (context.get("summary_rows") or [])
    }


def _meaningful(value: Any) -> str:
    text = str(value or "").strip()
    return "" if text in {"", EM_DASH, "None"} else text


def _text(value: Any) -> str:
    return _meaningful(value) or EM_DASH


def _testing_time(value: Any) -> str:
    try:
        days = int(value)
    except (TypeError, ValueError):
        return EM_DASH
    return f"{days} day" if days == 1 else f"{days} days"


def _humanize_action(value: Any) -> str:
    text = _meaningful(value)
    return text.replace("_", " ").title() if text else EM_DASH


def _rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)
