"""Services for the NABL AccredKit module."""

from __future__ import annotations

import csv
import io
import json
import re
import zipfile
from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime, timezone
from functools import lru_cache
from pathlib import Path
from typing import Any

from flask import current_app, url_for
from sqlalchemy import or_

from app.core.services.notifications import create_notification
from app.extensions import db
from app.models.core.notification import Notification
from app.models.core.user import User
from app.models.core.module_admin_assignment import ModuleAdminAssignment
from app.core.services.nabl_metadata import (
    DEFAULT_CHEMICALS,
    DEFAULT_RECORD_RESPONSIBLE,
    DEFAULT_RECORD_STORAGE,
    DOCUMENT_CATALOG,
    DOCUMENT_MAP,
    LAB_ROLE_OPTIONS,
    LABS,
    RETENTION_BY_SERIES,
    STANDARD_ANSWERS,
    STORAGE_OPTIONS,
    WIZARD_SECTION_MAP,
    WIZARD_SECTIONS,
)

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.shared import Inches, Pt
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:  # pragma: no cover
    Document = None
    WD_ALIGN_PARAGRAPH = None
    CT_Tbl = None
    CT_P = None
    Inches = None
    Pt = None
    Table = None
    Paragraph = None


BASE_DIR = Path(__file__).resolve().parents[3]
NABL_ROOT = BASE_DIR / "data" / "nabl"
NABL_BUILD_ROOT = NABL_ROOT / "builds"
IDWE_SOURCE_DIR = BASE_DIR / "NABL IDWE docs" / "IDWE Docs"
IDWE_MASTER_LIST_PATH = IDWE_SOURCE_DIR / "Procedures and Records List.docx"
LOGO_CANDIDATES = (
    BASE_DIR / "ONGC_Logo.jpeg",
    BASE_DIR / "ONGC Logo.jpeg",
    BASE_DIR / "ONGC_Logo.jpg",
    BASE_DIR / "ONGC Logo.jpg",
)

QUESTION_TO_PATH = {
    "Q1.2": ("q1", "test_parameters"),
    "Q1.3": ("q1", "measurement_units"),
    "Q1.4": ("q1", "test_methods"),
    "Q2A.2": ("q2", "organisation", "independence"),
    "Q2A.3": ("q2", "organisation", "roles"),
    "Q2B.1": ("q2", "personnel", "qualifications"),
    "Q2B.2": ("q2", "personnel", "competence_assessment"),
    "Q2B.3": ("q2", "personnel", "training"),
    "Q2B.4": ("q2", "personnel", "ongoing_monitoring"),
    "Q2C.1": ("q2", "facility", "description"),
    "Q2C.2": ("q2", "facility", "environmental_monitoring"),
    "Q2C.3": ("q2", "facility", "maintenance"),
    "Q2C.4": ("q2", "facility", "access_control"),
    "Q2D.2": ("q2", "equipment", "identification"),
    "Q2D.3": ("q2", "equipment", "calibration"),
    "Q2D.4": ("q2", "equipment", "crm"),
    "Q2D.5": ("q2", "equipment", "breakdown"),
    "Q2E.1": ("q2", "testing_process", "sample_receipt"),
    "Q2E.2": ("q2", "testing_process", "traceability"),
    "Q2E.3": ("q2", "testing_process", "nonconforming_samples"),
    "Q2E.4": ("q2", "testing_process", "contract_review"),
    "Q2F.1": ("q2", "qc", "frequency"),
    "Q2F.2": ("q2", "qc", "out_of_control"),
    "Q2F.3": ("q2", "qc", "proficiency_testing"),
    "Q2F.4": ("q2", "qc", "uncertainty_method"),
    "Q2G.1": ("q2", "reporting", "report_content"),
    "Q2G.2": ("q2", "reporting", "authorisation"),
    "Q2G.3": ("q2", "reporting", "amendments"),
    "Q2G.4": ("q2", "reporting", "turnaround"),
}

SECTION_TO_QUESTION = {row["slug"]: row["code"] for row in WIZARD_SECTIONS}

DOC_SOURCE_MAP = {
    "SOP-QS-01": "Q2A.2",
    "SOP-QS-02": "Q2C.4",
    "SOP-QS-03": "Q2B.1",
    "SOP-QS-04": "Q2C.1",
    "SOP-QS-05": "Q2D.2",
    "SOP-QS-06": "Q2D.3",
    "SOP-QS-08": "Q2E.4",
    "SOP-QS-09": "Q1.4",
    "SOP-QS-10": "Q2E.1",
    "SOP-QS-12": "Q2F.4",
    "SOP-QS-13": "Q2F.1",
    "SOP-QS-14": "Q2G.1",
}

DOC_STATE_PLACEHOLDER = "Placeholder"
DOC_STATE_NOT_INITIATED = "Not Initiated"
DOC_STATE_DRAFT_PUSHED = "Draft Pushed"
DOC_STATE_IN_DRAFTING = "In Drafting"
DOC_STATE_SUBMITTED = "Submitted"
DOC_STATE_RETURNED = "Returned"
DOC_STATE_PUBLISHED = "Published"

LAB_DRAFT_STATES = {
    DOC_STATE_DRAFT_PUSHED,
    DOC_STATE_IN_DRAFTING,
    DOC_STATE_RETURNED,
}

SOURCE_DOC_FILES = {
    "clause4": IDWE_SOURCE_DIR / "IDWE_Clause_4-1_4-2_Procedures_Records.docx",
    "clause5": IDWE_SOURCE_DIR / "IDWE_Clause_5_Structural_Requirements_Procedures_Records.docx",
    "clause6": IDWE_SOURCE_DIR / "IDWE_Clause6_Procedures_and_Records.docx",
    "clause7": IDWE_SOURCE_DIR / "IDWE_Clause7_Procedures_Records.docx",
}

SOURCE_BOUNDARY_RE = re.compile(r"^(?:ANNEX\s+[–-]\s+RECORD TEMPLATE\s+)?((?:SOP|QM|REC)-[A-Z0-9-]+)")


def ensure_storage_dirs() -> None:
    NABL_ROOT.mkdir(parents=True, exist_ok=True)
    NABL_BUILD_ROOT.mkdir(parents=True, exist_ok=True)
    (NABL_ROOT / "templates").mkdir(parents=True, exist_ok=True)


def _answers_path(lab_id: str) -> Path:
    return NABL_ROOT / f"{lab_id}_answers.json"


def _profile_path(lab_id: str) -> Path:
    return NABL_ROOT / f"{lab_id}_profile.json"


def _documents_path(lab_id: str) -> Path:
    return NABL_ROOT / f"{lab_id}_documents.json"


def _read_json(path: Path, default: Any) -> Any:
    if not path.exists():
        return deepcopy(default)
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def _write_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as handle:
        json.dump(payload, handle, indent=2, ensure_ascii=True)


def _utc_now_iso() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def _normalize_text(value: str) -> str:
    return " ".join((value or "").split()).strip()


def _normalized_doc_id_from_paragraph(text: str) -> str | None:
    clean = _normalize_text(text)
    if not clean:
        return None
    match = SOURCE_BOUNDARY_RE.match(clean)
    if not match:
        return None
    raw = match.group(1)
    if raw.startswith("REC-23") and len(raw) == 7:
        return raw
    return raw


@dataclass
class SimpleParagraph:
    text: str
    style_name: str = ""

    @property
    def style(self):
        return type("Style", (), {"name": self.style_name})()


def _iter_document_blocks(document: Document):
    body = document.element.body
    for child in body.iterchildren():
        if CT_P is not None and isinstance(child, CT_P):
            yield ("paragraph", Paragraph(child, document))
        elif CT_Tbl is not None and isinstance(child, CT_Tbl):
            yield ("table", Table(child, document))


def _paragraph_to_render_block(paragraph: Paragraph) -> dict[str, Any] | None:
    text = paragraph.text or ""
    style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    if not text.strip() and "Heading" not in style_name:
        return None
    block_type = "paragraph"
    if style_name.startswith("Heading"):
        block_type = "heading"
    elif style_name in {"Title", "Subtitle"}:
        block_type = "title"
    elif text.strip().startswith(("•", "-", "o", "▼")):
        block_type = "bullet"
    return {
        "type": block_type,
        "text": text,
        "style_name": style_name,
    }


def _table_to_render_block(table: Table) -> dict[str, Any]:
    rows = []
    max_cols = max((len(row.cells) for row in table.rows), default=0)
    for row in table.rows:
        rows.append([cell.text for cell in row.cells])
    return {
        "type": "table",
        "rows": rows,
        "cols": max_cols,
    }


@lru_cache(maxsize=1)
def get_idwe_master_catalog() -> dict[str, dict[str, Any]]:
    catalog: dict[str, dict[str, Any]] = {}
    if Document is None or not IDWE_MASTER_LIST_PATH.exists():
        return catalog

    document = Document(IDWE_MASTER_LIST_PATH)
    for table in document.tables:
        header = [_normalize_text(cell.text).lower() for cell in table.rows[0].cells]
        is_record_table = "record id" in header[0] if header else False
        for row in table.rows[1:]:
            values = [_normalize_text(cell.text) for cell in row.cells]
            if len(values) < 3:
                continue
            doc_id = values[0]
            if not doc_id:
                continue
            kind = "REC" if is_record_table or doc_id.startswith("REC-") else "SOP"
            catalog[doc_id] = {
                "doc_id": doc_id,
                "title": values[1],
                "clause": values[2],
                "kind": kind,
            }
    return catalog


def _source_file_for_doc(meta: dict[str, Any]) -> Path | None:
    clause = str(meta.get("clause", ""))
    if clause.startswith("4"):
        return SOURCE_DOC_FILES["clause4"]
    if clause.startswith("5"):
        return SOURCE_DOC_FILES["clause5"]
    if clause.startswith("6"):
        return SOURCE_DOC_FILES["clause6"]
    if clause.startswith("7"):
        return SOURCE_DOC_FILES["clause7"]
    return None


@lru_cache(maxsize=128)
def extract_idwe_reference(doc_id: str) -> dict[str, Any] | None:
    if Document is None:
        return None
    source_meta = get_idwe_master_catalog().get(doc_id) or DOCUMENT_MAP.get(doc_id)
    if not source_meta:
        return None
    source_file = _source_file_for_doc(source_meta)
    if source_file is None or not source_file.exists():
        return None

    document = Document(source_file)
    blocks = []
    for block_type, block in _iter_document_blocks(document):
        if block_type == "paragraph":
            text = _normalize_text(block.text)
            blocks.append({
                "type": "paragraph",
                "text": block.text,
                "normalized_text": text,
                "style_name": getattr(getattr(block, "style", None), "name", "") or "",
            })
        else:
            blocks.append({
                "type": "table",
                "rows": [[cell.text for cell in row.cells] for row in block.rows],
            })
    if not blocks:
        return None

    start_index = None
    for index, block in enumerate(blocks):
        if block["type"] != "paragraph":
            continue
        normalized_doc_id = _normalized_doc_id_from_paragraph(block["normalized_text"])
        if normalized_doc_id == doc_id:
            start_index = index
            break

    fallback_whole_file = False
    if start_index is None and doc_id == "SOP-STR-01":
        fallback_whole_file = True
        start_index = 0
    if start_index is None and doc_id == "SOP-QS-08":
        for index, block in enumerate(blocks):
            if block["type"] == "paragraph" and block["normalized_text"].startswith("1. Purpose"):
                start_index = index
                break
    if start_index is None:
        return {
            "doc_id": doc_id,
            "source_file": source_file.name,
            "title": source_meta.get("title", DOCUMENT_MAP.get(doc_id, {}).get("title", doc_id)),
            "clause": source_meta.get("clause", DOCUMENT_MAP.get(doc_id, {}).get("clause", "")),
            "found_exact_block": False,
            "fallback_whole_file": False,
            "blocks": [],
        }

    end_index = len(blocks)
    if not fallback_whole_file:
        for index in range(start_index + 1, len(blocks)):
            block = blocks[index]
            if block["type"] != "paragraph":
                continue
            normalized_doc_id = _normalized_doc_id_from_paragraph(block["normalized_text"])
            if normalized_doc_id and normalized_doc_id != doc_id:
                end_index = index
                break

    selected_blocks = blocks[start_index:end_index]
    return {
        "doc_id": doc_id,
        "source_file": source_file.name,
        "title": source_meta.get("title", DOCUMENT_MAP.get(doc_id, {}).get("title", doc_id)),
        "clause": source_meta.get("clause", DOCUMENT_MAP.get(doc_id, {}).get("clause", "")),
        "found_exact_block": not fallback_whole_file,
        "fallback_whole_file": fallback_whole_file,
        "blocks": selected_blocks,
    }


def build_idwe_reference_view(doc_id: str) -> dict[str, Any] | None:
    reference = extract_idwe_reference(doc_id)
    if not reference:
        return None
    rendered_blocks: list[dict[str, Any]] = []
    for block in reference["blocks"]:
        if block["type"] == "paragraph":
            paragraph = _paragraph_to_render_block(
                SimpleParagraph(block["text"], block.get("style_name", ""))
            )
            if paragraph is not None:
                rendered_blocks.append(paragraph)
            continue
        rendered_blocks.append(
            {
                "type": "table",
                "rows": block["rows"],
                "cols": max((len(row) for row in block["rows"]), default=0),
            }
        )
    return {
        "doc_id": reference["doc_id"],
        "source_file": reference["source_file"],
        "title": reference["title"],
        "clause": reference["clause"],
        "found_exact_block": reference["found_exact_block"],
        "fallback_whole_file": reference["fallback_whole_file"],
        "blocks": rendered_blocks,
    }


def _document_to_render_blocks(document: Document) -> list[dict[str, Any]]:
    rendered_blocks: list[dict[str, Any]] = []
    for block_type, block in _iter_document_blocks(document):
        if block_type == "paragraph":
            paragraph = _paragraph_to_render_block(block)
            if paragraph is not None:
                rendered_blocks.append(paragraph)
            continue
        rendered_blocks.append(_table_to_render_block(block))
    return rendered_blocks


def build_generated_document_view(
    lab_id: str,
    doc_id: str,
    *,
    published_version: str | None = None,
) -> dict[str, Any] | None:
    if Document is None:  # pragma: no cover
        return None
    docx_bytes = build_docx_bytes(lab_id, doc_id, published_version=published_version)
    document = Document(io.BytesIO(docx_bytes))
    meta = get_resolved_document_meta(doc_id)
    return {
        "doc_id": doc_id,
        "source_file": meta.get("source_file"),
        "title": meta["title"],
        "clause": meta["clause"],
        "published_version": published_version,
        "blocks": _document_to_render_blocks(document),
    }


def get_resolved_document_meta(doc_id: str) -> dict[str, Any]:
    base = deepcopy(DOCUMENT_MAP[doc_id])
    source_meta = get_idwe_master_catalog().get(doc_id)
    if source_meta:
        base["title"] = source_meta.get("title") or base["title"]
        base["clause"] = source_meta.get("clause") or base["clause"]
    reference = extract_idwe_reference(doc_id)
    base["source_file"] = reference["source_file"] if reference else None
    base["source_available"] = reference is not None
    return base


def get_source_document_summary() -> dict[str, Any]:
    catalog = get_idwe_master_catalog()
    available_files = [path for path in SOURCE_DOC_FILES.values() if path.exists()]
    return {
        "available": bool(available_files),
        "files_count": len(available_files) + (1 if IDWE_MASTER_LIST_PATH.exists() else 0),
        "catalog_count": len(catalog),
        "source_dir": str(IDWE_SOURCE_DIR.relative_to(BASE_DIR)) if IDWE_SOURCE_DIR.exists() else None,
        "files": [path.name for path in available_files] + ([IDWE_MASTER_LIST_PATH.name] if IDWE_MASTER_LIST_PATH.exists() else []),
    }


def get_lab_definition(lab_id: str) -> dict[str, Any] | None:
    return deepcopy(LABS.get(lab_id))


def get_all_labs() -> list[dict[str, Any]]:
    return [deepcopy(row) for row in LABS.values()]


def _default_chemicals(lab_id: str) -> list[dict[str, str]]:
    return [
        {"name": name, "spec": spec}
        for name, spec in DEFAULT_CHEMICALS.get(lab_id, [])
    ]


def default_profile(lab_id: str) -> dict[str, Any]:
    definition = LABS[lab_id]
    return {
        "lab_id": lab_id,
        "display_name": definition["display_name"],
        "location": definition["location"],
        "parent_organisation": definition["parent_organisation"],
        "approved_by": definition["approved_by"],
        "lab_code": definition["lab_code"],
        "version": 1,
        "updated_at": _utc_now_iso(),
    }


def default_answers(lab_id: str) -> dict[str, Any]:
    return {
        "lab_id": lab_id,
        "q1": {
            "chemicals": _default_chemicals(lab_id),
            "test_parameters": {"choice": "standard", "custom_text": ""},
            "measurement_units": {"choice": "standard", "custom_text": ""},
            "test_methods": {"choice": "standard", "custom_text": ""},
        },
        "q2": {
            "organisation": {
                "independence": {"choice": "standard", "custom_text": ""},
                "roles": {"choice": "standard", "custom_text": ""},
                "approval_authority": LABS[lab_id]["approved_by"],
            },
            "personnel": {
                "qualifications": {"choice": "standard", "custom_text": ""},
                "competence_assessment": {"choice": "standard", "custom_text": ""},
                "training": {"choice": "standard", "custom_text": ""},
                "ongoing_monitoring": {"choice": "standard", "custom_text": ""},
            },
            "facility": {
                "description": {"choice": "standard", "custom_text": ""},
                "environmental_monitoring": {"choice": "standard", "custom_text": ""},
                "maintenance": {"choice": "standard", "custom_text": ""},
                "access_control": {"choice": "standard", "custom_text": ""},
            },
            "equipment": {
                "entry_method": "manual",
                "equipment_list": [],
                "identification": {"choice": "standard", "custom_text": ""},
                "calibration": {"choice": "standard", "custom_text": ""},
                "crm": {"choice": "standard", "custom_text": ""},
                "breakdown": {"choice": "standard", "custom_text": ""},
            },
            "testing_process": {
                "sample_receipt": {"choice": "standard", "custom_text": ""},
                "traceability": {"choice": "standard", "custom_text": ""},
                "nonconforming_samples": {"choice": "standard", "custom_text": ""},
                "contract_review": {"choice": "standard", "custom_text": ""},
                "lsn_format": "[LAB]-[YEAR]-[NNN]",
            },
            "qc": {
                "frequency": {"choice": "standard", "custom_text": ""},
                "out_of_control": {"choice": "standard", "custom_text": ""},
                "proficiency_testing": {"choice": "standard", "custom_text": ""},
                "uncertainty_method": {"choice": "standard", "custom_text": ""},
            },
            "reporting": {
                "report_content": {"choice": "standard", "custom_text": ""},
                "authorisation": {"choice": "standard", "custom_text": ""},
                "amendments": {"choice": "standard", "custom_text": ""},
                "turnaround": {"choice": "standard", "custom_text": ""},
                "report_number_format": "[LAB]-TR-[YEAR]-[NNN]",
            },
        },
        "q3": {"records": []},
        "wizard_progress": {
            "saved_sections": [],
            "last_section": None,
        },
        "sop_sections": {},
        "rec_sections": {},
        "last_saved": _utc_now_iso(),
        "wizard_complete": False,
    }


def _default_document_row(meta: dict[str, Any]) -> dict[str, Any]:
    base_state = (
        DOC_STATE_PLACEHOLDER
        if meta["layer"] not in {"layer1", "placeholder"}
        else DOC_STATE_NOT_INITIATED
    )
    return {
        "doc_id": meta["doc_id"],
        "title": meta["title"],
        "clause": meta["clause"],
        "kind": meta["kind"],
        "layer": meta["layer"],
        "compile": bool(meta["compile"]),
        "state": base_state,
        "custom_sections_count": 0,
        "draft_pushed_at": None,
        "draft_pushed_by": None,
        "reviewed_at": None,
        "submitted_at": None,
        "submitted_by": None,
        "returned_at": None,
        "returned_by": None,
        "return_comments": "",
        "return_comments_by_section": {},
        "current_draft": {
            "sections": [],
            "updated_at": _utc_now_iso(),
            "updated_by": "system",
        },
        "versions": [],
    }


def default_documents(lab_id: str) -> dict[str, Any]:
    return {
        "lab_id": lab_id,
        "documents": {
            meta["doc_id"]: _default_document_row(meta)
            for meta in DOCUMENT_CATALOG
        },
        "updated_at": _utc_now_iso(),
    }


def load_lab_profile(lab_id: str) -> dict[str, Any]:
    ensure_storage_dirs()
    payload = _read_json(_profile_path(lab_id), default_profile(lab_id))
    merged = default_profile(lab_id)
    merged.update(payload)
    return merged


def save_lab_profile(lab_id: str, profile: dict[str, Any]) -> dict[str, Any]:
    ensure_storage_dirs()
    payload = default_profile(lab_id)
    payload.update(profile or {})
    payload["updated_at"] = _utc_now_iso()
    _write_json(_profile_path(lab_id), payload)
    return payload


def load_lab_answers(lab_id: str) -> dict[str, Any]:
    ensure_storage_dirs()
    payload = _read_json(_answers_path(lab_id), default_answers(lab_id))
    merged = default_answers(lab_id)
    merged = _deep_merge(merged, payload)
    if not merged["q3"]["records"]:
        merged["q3"]["records"] = build_default_records_rows()
    return merged


def save_lab_answers(lab_id: str, answers: dict[str, Any]) -> dict[str, Any]:
    ensure_storage_dirs()
    answers["last_saved"] = _utc_now_iso()
    _write_json(_answers_path(lab_id), answers)
    return answers


def load_lab_documents(lab_id: str) -> dict[str, Any]:
    ensure_storage_dirs()
    payload = _read_json(_documents_path(lab_id), default_documents(lab_id))
    merged = default_documents(lab_id)
    merged = _deep_merge(merged, payload)
    return merged


def save_lab_documents(lab_id: str, documents: dict[str, Any]) -> dict[str, Any]:
    ensure_storage_dirs()
    documents["updated_at"] = _utc_now_iso()
    _write_json(_documents_path(lab_id), documents)
    return documents


def ensure_lab_initialized(lab_id: str) -> tuple[dict[str, Any], dict[str, Any], dict[str, Any]]:
    profile = load_lab_profile(lab_id)
    answers = load_lab_answers(lab_id)
    documents = load_lab_documents(lab_id)
    sync_lab_documents(lab_id, answers=answers, profile=profile, documents=documents)
    answers = load_lab_answers(lab_id)
    documents = load_lab_documents(lab_id)
    return profile, answers, documents


def _deep_merge(base: Any, update: Any) -> Any:
    if isinstance(base, dict) and isinstance(update, dict):
        merged = deepcopy(base)
        for key, value in update.items():
            if key in merged:
                merged[key] = _deep_merge(merged[key], value)
            else:
                merged[key] = deepcopy(value)
        return merged
    if isinstance(base, list) and isinstance(update, list):
        return deepcopy(update)
    return deepcopy(update)


def get_standard_answer(question_code: str, profile: dict[str, Any]) -> str:
    template = STANDARD_ANSWERS.get(question_code, "")
    return (
        template
        .replace("[Lab Name]", profile["display_name"])
        .replace("[Location]", profile["location"])
        .replace("[LAB]", profile["lab_code"])
    )


def _get_nested(payload: dict[str, Any], path: tuple[str, ...]) -> Any:
    cursor = payload
    for part in path:
        cursor = cursor[part]
    return cursor


def _set_nested(payload: dict[str, Any], path: tuple[str, ...], value: Any) -> None:
    cursor = payload
    for part in path[:-1]:
        cursor = cursor[part]
    cursor[path[-1]] = value


def get_answer_text(question_code: str, answers: dict[str, Any], profile: dict[str, Any]) -> str:
    if question_code == "Q2A.1":
        return (
            f"{profile['display_name']}, {profile['location']}. "
            f"Parent organisation: {profile['parent_organisation']}."
        )
    if question_code == "Q2A.4":
        return profile["approved_by"]
    if question_code == "Q1.1":
        chemicals = answers["q1"]["chemicals"]
        return "\n".join(
            f"- {row['name']} ({row['spec']})"
            for row in chemicals
            if row.get("name")
        ) or "No chemicals configured."

    path = QUESTION_TO_PATH.get(question_code)
    if not path:
        return ""
    answer_row = _get_nested(answers, path)
    if isinstance(answer_row, dict) and "choice" in answer_row:
        if answer_row.get("choice") == "own":
            return (answer_row.get("custom_text") or "").strip()
        return get_standard_answer(question_code, profile)
    if isinstance(answer_row, str):
        return answer_row
    return ""


def _choice_payload_from_request(question_code: str, payload: dict[str, Any], profile: dict[str, Any]) -> dict[str, str]:
    choice = (payload.get("choice") or "standard").strip().lower()
    custom_text = (payload.get("custom_text") or "").strip()
    if choice not in {"standard", "own"}:
        choice = "standard"
    if choice == "own" and not custom_text:
        custom_text = get_standard_answer(question_code, profile)
    return {"choice": choice, "custom_text": custom_text}


def build_default_records_rows() -> list[dict[str, Any]]:
    rows = []
    for meta in DOCUMENT_CATALOG:
        if meta["kind"] != "REC":
            continue
        rows.append(
            {
                "rec_id": meta["doc_id"],
                "title": meta["title"],
                "responsible": DEFAULT_RECORD_RESPONSIBLE,
                "retention": retention_for_record(meta["doc_id"]),
                "storage": DEFAULT_RECORD_STORAGE,
                "confirmed": meta["layer"] == "layer1",
                "available": meta["layer"] == "layer1",
            }
        )
    return rows


def retention_for_record(rec_id: str) -> str:
    if rec_id in RETENTION_BY_SERIES:
        return RETENTION_BY_SERIES[rec_id]
    prefix = rec_id.split("-")[0] + "-" + rec_id.split("-")[1]
    return RETENTION_BY_SERIES.get(prefix, "5 years")


def progress_for_answers(answers: dict[str, Any]) -> dict[str, Any]:
    saved = set(answers.get("wizard_progress", {}).get("saved_sections", []))
    total = len(WIZARD_SECTIONS)
    completed = len(saved)
    by_step = {}
    for section in WIZARD_SECTIONS:
        step = section["step"]
        bucket = by_step.setdefault(step, {"completed": 0, "total": 0})
        bucket["total"] += 1
        if section["slug"] in saved:
            bucket["completed"] += 1
    return {
        "completed": completed,
        "total": total,
        "percent": int(round((completed / total) * 100)) if total else 0,
        "steps": by_step,
    }


def section_navigation(section_slug: str) -> dict[str, Any]:
    slugs = [row["slug"] for row in WIZARD_SECTIONS]
    index = slugs.index(section_slug)
    previous_slug = slugs[index - 1] if index > 0 else None
    next_slug = slugs[index + 1] if index < len(slugs) - 1 else None
    return {"previous": previous_slug, "next": next_slug}


def save_wizard_section(
    lab_id: str,
    section_slug: str,
    payload: dict[str, Any],
    files: dict[str, Any] | None = None,
) -> tuple[dict[str, Any], dict[str, Any], dict[str, Any]]:
    profile, answers, documents = ensure_lab_initialized(lab_id)
    section = WIZARD_SECTION_MAP[section_slug]
    question_code = section["code"]

    if section["type"] == "chemical_table":
        chemicals = payload.get("chemicals")
        if chemicals is None:
            names = payload.get("chemical_name", [])
            specs = payload.get("chemical_spec", [])
            chemicals = [
                {"name": name.strip(), "spec": (specs[index] if index < len(specs) else "").strip()}
                for index, name in enumerate(names)
                if name.strip()
            ]
        answers["q1"]["chemicals"] = chemicals or _default_chemicals(lab_id)

    elif section["type"] == "lab_profile":
        profile = save_lab_profile(
            lab_id,
            {
                "display_name": (payload.get("display_name") or profile["display_name"]).strip(),
                "location": (payload.get("location") or profile["location"]).strip(),
                "parent_organisation": (
                    payload.get("parent_organisation") or profile["parent_organisation"]
                ).strip(),
            },
        )

    elif section["type"] == "approval_authority":
        approval_authority = (payload.get("approval_authority") or profile["approved_by"]).strip()
        profile = save_lab_profile(lab_id, {"approved_by": approval_authority})
        answers["q2"]["organisation"]["approval_authority"] = approval_authority

    elif section["type"] == "equipment_entry":
        entry_method = (payload.get("entry_method") or "manual").strip().lower()
        if entry_method not in {"upload", "manual"}:
            entry_method = "manual"
        equipment_list = []
        upload_file = None
        if files:
            upload_file = files.get("equipment_file")
        if entry_method == "upload" and upload_file and getattr(upload_file, "filename", ""):
            equipment_list = parse_equipment_upload(upload_file.filename, upload_file.read())
        else:
            names = payload.get("equipment_name", [])
            makes = payload.get("equipment_make", [])
            models = payload.get("equipment_model", [])
            serials = payload.get("equipment_serial", [])
            dues = payload.get("equipment_cal_due", [])
            agencies = payload.get("equipment_cal_agency", [])
            for index, name in enumerate(names):
                clean_name = (name or "").strip()
                if not clean_name:
                    continue
                equipment_list.append(
                    {
                        "name": clean_name,
                        "make": (makes[index] if index < len(makes) else "").strip(),
                        "model": (models[index] if index < len(models) else "").strip(),
                        "serial": (serials[index] if index < len(serials) else "").strip(),
                        "cal_due": (dues[index] if index < len(dues) else "").strip(),
                        "cal_agency": (agencies[index] if index < len(agencies) else "").strip(),
                    }
                )
        answers["q2"]["equipment"]["entry_method"] = entry_method
        answers["q2"]["equipment"]["equipment_list"] = equipment_list

    elif section["type"] == "records_table":
        rows = []
        rec_ids = payload.get("rec_id", [])
        if rec_ids:
            confirmed_ids = set(payload.get("confirmed", []))
            titles = payload.get("title", [])
            responsible = payload.get("responsible", [])
            retention = payload.get("retention", [])
            storage = payload.get("storage", [])
            for index, rec_id in enumerate(rec_ids):
                rows.append(
                    {
                        "rec_id": rec_id,
                        "title": titles[index] if index < len(titles) else DOCUMENT_MAP.get(rec_id, {}).get("title", rec_id),
                        "responsible": responsible[index] if index < len(responsible) else DEFAULT_RECORD_RESPONSIBLE,
                        "retention": retention[index] if index < len(retention) else retention_for_record(rec_id),
                        "storage": storage[index] if index < len(storage) else DEFAULT_RECORD_STORAGE,
                        "confirmed": rec_id in confirmed_ids,
                        "available": DOCUMENT_MAP.get(rec_id, {}).get("layer") == "layer1",
                    }
                )
        else:
            rows = payload.get("records") or build_default_records_rows()
        answers["q3"]["records"] = rows

    else:
        choice_value = _choice_payload_from_request(question_code, payload, profile)
        answer_path = QUESTION_TO_PATH[question_code]
        _set_nested(answers, answer_path, choice_value)

    progress_state = answers.setdefault("wizard_progress", {"saved_sections": [], "last_section": None})
    if section_slug not in progress_state["saved_sections"]:
        progress_state["saved_sections"].append(section_slug)
    progress_state["last_section"] = section_slug
    answers["wizard_complete"] = len(progress_state["saved_sections"]) == len(WIZARD_SECTIONS)
    save_lab_answers(lab_id, answers)
    sync_lab_documents(lab_id, answers=answers, profile=profile, documents=documents)
    return profile, load_lab_answers(lab_id), load_lab_documents(lab_id)


def parse_equipment_upload(filename: str, file_bytes: bytes) -> list[dict[str, str]]:
    filename = (filename or "").lower()
    if filename.endswith(".csv"):
        decoded = file_bytes.decode("utf-8-sig")
        reader = csv.DictReader(io.StringIO(decoded))
        rows = list(reader)
    else:
        try:
            import pandas as pd
        except ImportError as exc:  # pragma: no cover
            raise ValueError("Equipment upload requires pandas for Excel files.") from exc
        dataframe = pd.read_excel(io.BytesIO(file_bytes), dtype=str).fillna("")
        rows = dataframe.to_dict(orient="records")

    normalized = []
    for row in rows:
        lookup = {re.sub(r"[^a-z0-9]+", "", str(key).lower()): str(value).strip() for key, value in row.items()}
        name = lookup.get("equipmentname") or lookup.get("name")
        if not name:
            continue
        normalized.append(
            {
                "name": name,
                "make": lookup.get("make", ""),
                "model": lookup.get("model", ""),
                "serial": lookup.get("serialno") or lookup.get("serialnumber") or lookup.get("serial", ""),
                "cal_due": lookup.get("calibrationduedate") or lookup.get("calduedate") or lookup.get("caldue", ""),
                "cal_agency": lookup.get("calibratingagency") or lookup.get("calibrationagency") or lookup.get("agency", ""),
            }
        )
    return normalized


def _lab_context(lab_id: str, answers: dict[str, Any], profile: dict[str, Any]) -> dict[str, Any]:
    chemicals = answers["q1"]["chemicals"]
    equipment = answers["q2"]["equipment"]["equipment_list"]
    record_rows = answers["q3"]["records"] or build_default_records_rows()
    return {
        "lab_id": lab_id,
        "profile": profile,
        "answers": answers,
        "chemicals": chemicals,
        "equipment": equipment,
        "records": record_rows,
        "scope_summary": ", ".join(
            f"{row['name']} ({row['spec']})" if row.get("spec") else row["name"]
            for row in chemicals
            if row.get("name")
        ) or "Scope pending confirmation.",
        "equipment_id_format": f"{profile['lab_code']}-EQ-[NNN]",
        "lsn_format": f"{profile['lab_code']}-[YEAR]-[NNN]",
        "report_number_format": f"{profile['lab_code']}-TR-[YEAR]-[NNN]",
    }


def _make_section(
    section_id: str,
    title: str,
    body: str,
    *,
    source_question: str | None = None,
    answers: dict[str, Any] | None = None,
    profile: dict[str, Any] | None = None,
) -> dict[str, Any]:
    choice = "standard"
    custom_text = ""
    if source_question and answers and profile:
        path = QUESTION_TO_PATH.get(source_question)
        if path:
            payload = _get_nested(answers, path)
            if isinstance(payload, dict):
                choice = payload.get("choice", "standard")
                custom_text = payload.get("custom_text", "")
    return {
        "section_id": section_id,
        "title": title,
        "body": body.strip(),
        "auto_body": body.strip(),
        "choice": choice,
        "custom_text": custom_text,
        "source_question": source_question,
        "auto_filled": True,
        "reviewed": False,
        "manual_override": False,
    }


def build_document_sections(doc_id: str, context: dict[str, Any]) -> list[dict[str, Any]]:
    answers = context["answers"]
    profile = context["profile"]
    meta = get_resolved_document_meta(doc_id)
    if meta["layer"] == "placeholder":
        return [
            _make_section(
                "section_placeholder",
                "Clause 8 Placeholder",
                "This document is included as a Clause 8 placeholder in Layer 1. "
                "Detailed controlled text can be loaded from the IDWE template once "
                "the final governed source document is available.",
            )
        ]
    if meta["layer"] in {"layer2", "layer3"}:
        return [
            _make_section(
                "section_future",
                "Future Build Layer",
                "This document is tracked in the registry but is not part of the "
                "current Layer 1 generation scope.",
            )
        ]

    common_reference = (
        f"Laboratory: {profile['display_name']} | Location: {profile['location']} | "
        f"Approved by: {profile['approved_by']}"
    )
    scope_summary = context["scope_summary"]
    q1_methods = get_answer_text("Q1.4", answers, profile)
    q2_roles = get_answer_text("Q2A.3", answers, profile)
    q2_identity = get_answer_text("Q2A.1", answers, profile)
    q2_reporting = get_answer_text("Q2G.2", answers, profile)
    q2_receipt = get_answer_text("Q2E.1", answers, profile)

    if doc_id == "SOP-STR-01":
        return [
            _make_section("section_1", "Organisation Context", q2_identity, source_question="Q2A.1", answers=answers, profile=profile),
            _make_section("section_2", "Roles and Reporting", q2_roles, source_question="Q2A.3", answers=answers, profile=profile),
            _make_section("section_3", "Approval Authority", profile["approved_by"], source_question="Q2A.4", answers=answers, profile=profile),
            _make_section("section_4", "Scope Reference", f"Current confirmed chemical scope:\n{scope_summary}", source_question="Q1.1", answers=answers, profile=profile),
        ]
    if doc_id == "QM-01":
        return [
            _make_section("section_1", "Quality Policy Context", common_reference),
            _make_section("section_2", "Scope of Accreditation Readiness", f"Confirmed chemical scope:\n{scope_summary}", source_question="Q1.1", answers=answers, profile=profile),
            _make_section("section_3", "Testing Standards", q1_methods, source_question="Q1.4", answers=answers, profile=profile),
            _make_section("section_4", "Reporting Governance", q2_reporting, source_question="Q2G.2", answers=answers, profile=profile),
        ]
    if doc_id == "SOP-QS-08":
        return [
            _make_section("section_5_1", "Scope Table", f"The laboratory currently confirms the following chemical scope:\n{scope_summary}", source_question="Q1.1", answers=answers, profile=profile),
            _make_section("section_5_2", "Request Review Workflow", get_answer_text("Q2E.4", answers, profile), source_question="Q2E.4", answers=answers, profile=profile),
            _make_section("section_5_3", "Sample Intake Reference", q2_receipt, source_question="Q2E.1", answers=answers, profile=profile),
        ]
    if doc_id == "SOP-QS-09":
        return [
            _make_section("section_5_1", "Applicable Standards", q1_methods, source_question="Q1.4", answers=answers, profile=profile),
            _make_section("section_5_2", "Measurement Units and Ranges", get_answer_text("Q1.3", answers, profile), source_question="Q1.3", answers=answers, profile=profile),
            _make_section("section_5_3", "Method Coverage", get_answer_text("Q1.2", answers, profile), source_question="Q1.2", answers=answers, profile=profile),
        ]
    if doc_id == "SOP-QS-10":
        return [
            _make_section("section_5_1", "Receipt and Registration", q2_receipt, source_question="Q2E.1", answers=answers, profile=profile),
            _make_section("section_5_2", "Traceability Controls", get_answer_text("Q2E.2", answers, profile), source_question="Q2E.2", answers=answers, profile=profile),
            _make_section("section_5_3", "Non-conforming Samples", get_answer_text("Q2E.3", answers, profile), source_question="Q2E.3", answers=answers, profile=profile),
        ]
    if doc_id == "SOP-QS-12":
        return [
            _make_section("section_5_1", "Uncertainty Method", get_answer_text("Q2F.4", answers, profile), source_question="Q2F.4", answers=answers, profile=profile),
            _make_section("section_5_2", "Scope Reference", get_answer_text("Q1.2", answers, profile), source_question="Q1.2", answers=answers, profile=profile),
        ]
    if doc_id == "SOP-QS-13":
        return [
            _make_section("section_5_1", "QC Frequency", get_answer_text("Q2F.1", answers, profile), source_question="Q2F.1", answers=answers, profile=profile),
            _make_section("section_5_2", "Out-of-Control Response", get_answer_text("Q2F.2", answers, profile), source_question="Q2F.2", answers=answers, profile=profile),
            _make_section("section_5_3", "Proficiency Testing", get_answer_text("Q2F.3", answers, profile), source_question="Q2F.3", answers=answers, profile=profile),
        ]
    if doc_id == "SOP-QS-14":
        return [
            _make_section("section_5_1", "Report Content", get_answer_text("Q2G.1", answers, profile), source_question="Q2G.1", answers=answers, profile=profile),
            _make_section("section_5_2", "Review and Authorisation", q2_reporting, source_question="Q2G.2", answers=answers, profile=profile),
            _make_section("section_5_3", "Amendments", get_answer_text("Q2G.3", answers, profile), source_question="Q2G.3", answers=answers, profile=profile),
            _make_section("section_5_4", "Turnaround Commitment", get_answer_text("Q2G.4", answers, profile), source_question="Q2G.4", answers=answers, profile=profile),
        ]

    source_question = DOC_SOURCE_MAP.get(doc_id)
    configured_body = get_answer_text(source_question, answers, profile) if source_question else common_reference
    generic_sections = [
        _make_section("section_1", "Document Context", common_reference),
        _make_section("section_2", "Configured Laboratory Input", configured_body, source_question=source_question, answers=answers, profile=profile),
        _make_section(
            "section_3",
            "Referenced Scope and Records",
            f"Chemical scope summary:\n{scope_summary}\n\nRelated roles:\n{q2_roles}",
            source_question="Q1.1",
            answers=answers,
            profile=profile,
        ),
    ]
    if doc_id.startswith("REC-"):
        return build_record_sections(doc_id, context)
    return generic_sections


def build_record_sections(doc_id: str, context: dict[str, Any]) -> list[dict[str, Any]]:
    answers = context["answers"]
    profile = context["profile"]
    meta = DOCUMENT_MAP[doc_id]

    if doc_id == "REC-04A":
        return [
            _make_section("section_1", "Organisation Header", get_answer_text("Q2A.1", answers, profile), source_question="Q2A.1", answers=answers, profile=profile),
            _make_section("section_2", "Reporting Structure", get_answer_text("Q2A.3", answers, profile), source_question="Q2A.3", answers=answers, profile=profile),
        ]
    if doc_id == "REC-04C":
        return [
            _make_section("section_1", "Role Catalogue", "\n".join(f"- {role}" for role in LAB_ROLE_OPTIONS)),
            _make_section("section_2", "Reporting Lines", get_answer_text("Q2A.3", answers, profile), source_question="Q2A.3", answers=answers, profile=profile),
        ]
    if doc_id in {"REC-09A", "REC-09B"}:
        equipment_lines = [
            f"- {row['name']} | {row.get('make', '')} | {row.get('model', '')} | "
            f"{row.get('serial', '')} | Due: {row.get('cal_due', '')} | Agency: {row.get('cal_agency', '')}"
            for row in context["equipment"]
        ] or ["- No equipment entered yet."]
        return [
            _make_section("section_1", "Equipment ID Format", f"Configured ID format: {context['equipment_id_format']}", source_question="Q2D.2", answers=answers, profile=profile),
            _make_section("section_2", "Configured Equipment Rows", "\n".join(equipment_lines)),
            _make_section("section_3", "Calibration Control", get_answer_text("Q2D.3", answers, profile), source_question="Q2D.3", answers=answers, profile=profile),
        ]
    if doc_id == "REC-17A":
        return [
            _make_section("section_1", "Contract Review Prerequisite", get_answer_text("Q2E.4", answers, profile), source_question="Q2E.4", answers=answers, profile=profile),
            _make_section("section_2", "Requested Scope", context["scope_summary"], source_question="Q1.1", answers=answers, profile=profile),
        ]
    if doc_id in {"REC-22A", "REC-22B"}:
        return [
            _make_section("section_1", "Receipt Workflow", get_answer_text("Q2E.1", answers, profile), source_question="Q2E.1", answers=answers, profile=profile),
            _make_section("section_2", "LSN Format", f"Configured LSN format: {context['lsn_format']}"),
        ]
    if doc_id == "REC-32A":
        lines = [
            f"- {meta_row['doc_id']} | {meta_row['title']} | {meta_row['layer']}"
            for meta_row in DOCUMENT_CATALOG
            if meta_row["kind"] in {"SOP", "REC"}
        ]
        return [
            _make_section("section_1", "Generated Master List", "\n".join(lines)),
            _make_section("section_2", "Governance Context", f"Approved by: {profile['approved_by']}"),
        ]
    if doc_id == "REC-33A":
        lines = [
            f"- {row['rec_id']} | {row['responsible']} | {row['retention']} | {row['storage']}"
            for row in answers["q3"]["records"]
        ]
        return [
            _make_section("section_1", "Retention Schedule", "\n".join(lines), source_question="Q3", answers=answers, profile=profile),
        ]
    if doc_id == "REC-07A":
        return [
            _make_section("section_1", "Environmental Limits", get_answer_text("Q2C.2", answers, profile), source_question="Q2C.2", answers=answers, profile=profile),
        ]
    if doc_id in {"REC-28B", "REC-28C", "REC-28D"}:
        return [
            _make_section("section_1", "Reporting Governance", get_answer_text("Q2G.2", answers, profile), source_question="Q2G.2", answers=answers, profile=profile),
            _make_section("section_2", "Report Number Format", f"Configured report number format: {context['report_number_format']}"),
        ]
    if doc_id in {"REC-01B", "REC-29A", "REC-29B", "REC-30A", "REC-30B", "REC-30C", "REC-36A", "REC-36B", "REC-36C"}:
        source_question = "Q2A.2" if doc_id == "REC-01B" else "Q2F.2" if doc_id.startswith("REC-36") else None
        body = get_answer_text(source_question, answers, profile) if source_question else f"{meta['title']} configured for {profile['display_name']}."
        return [
            _make_section("section_1", "Record Purpose", body, source_question=source_question, answers=answers, profile=profile),
            _make_section("section_2", "Ownership", f"Responsible role default: {DEFAULT_RECORD_RESPONSIBLE}"),
        ]

    return [
        _make_section("section_1", "Record Purpose", f"{meta['title']} configured for {profile['display_name']}."),
        _make_section("section_2", "Laboratory Reference", f"Scope summary:\n{context['scope_summary']}"),
    ]


def _update_section_registry(answers: dict[str, Any], doc_id: str, sections: list[dict[str, Any]]) -> None:
    target = answers["sop_sections"] if doc_id.startswith(("SOP", "QM")) else answers["rec_sections"]
    target[doc_id] = {}
    for section in sections:
        target[doc_id][section["section_id"]] = {
            "choice": section["choice"],
            "custom_text": section["custom_text"],
            "source_question": section["source_question"],
            "auto_filled": section["auto_filled"],
            "reviewed": section["reviewed"],
        }


def sync_lab_documents(
    lab_id: str,
    *,
    answers: dict[str, Any] | None = None,
    profile: dict[str, Any] | None = None,
    documents: dict[str, Any] | None = None,
) -> dict[str, Any]:
    profile = profile or load_lab_profile(lab_id)
    answers = answers or load_lab_answers(lab_id)
    documents = documents or load_lab_documents(lab_id)
    if not answers["q3"]["records"]:
        answers["q3"]["records"] = build_default_records_rows()

    context = _lab_context(lab_id, answers, profile)
    for meta in DOCUMENT_CATALOG:
        doc_id = meta["doc_id"]
        row = documents["documents"].setdefault(doc_id, _default_document_row(meta))
        generated_sections = build_document_sections(doc_id, context)
        existing_sections = {
            section["section_id"]: section
            for section in row["current_draft"].get("sections", [])
        }
        merged_sections = []
        for generated in generated_sections:
            existing = existing_sections.get(generated["section_id"])
            if existing and existing.get("manual_override"):
                generated["body"] = existing.get("body", generated["body"])
                generated["choice"] = "custom"
                generated["manual_override"] = True
            generated["reviewed"] = bool(existing.get("reviewed")) if existing else False
            merged_sections.append(generated)
        row["current_draft"]["sections"] = merged_sections
        row["current_draft"]["updated_at"] = _utc_now_iso()
        row["custom_sections_count"] = sum(
            1
            for section in merged_sections
            if section.get("manual_override") or section.get("choice") == "own"
        )
        if row["state"] not in {
            DOC_STATE_NOT_INITIATED,
            DOC_STATE_DRAFT_PUSHED,
            DOC_STATE_IN_DRAFTING,
            DOC_STATE_SUBMITTED,
            DOC_STATE_RETURNED,
            DOC_STATE_PUBLISHED,
        }:
            row["state"] = DOC_STATE_PLACEHOLDER if meta["layer"] not in {"layer1", "placeholder"} else DOC_STATE_NOT_INITIATED
        if meta["layer"] == "placeholder":
            row["state"] = row["state"] if row["state"] != DOC_STATE_PLACEHOLDER else DOC_STATE_PLACEHOLDER
        if meta["layer"] in {"layer2", "layer3"}:
            row["state"] = DOC_STATE_PLACEHOLDER
        _update_section_registry(answers, doc_id, merged_sections)

    save_lab_answers(lab_id, answers)
    return save_lab_documents(lab_id, documents)


def get_document_list(lab_id: str) -> list[dict[str, Any]]:
    _profile, _answers, documents = ensure_lab_initialized(lab_id)
    rows = []
    for meta in DOCUMENT_CATALOG:
        resolved_meta = get_resolved_document_meta(meta["doc_id"])
        row = documents["documents"][meta["doc_id"]]
        rows.append(
            {
                "doc_id": resolved_meta["doc_id"],
                "title": resolved_meta["title"],
                "clause": resolved_meta["clause"],
                "kind": resolved_meta["kind"],
                "layer": resolved_meta["layer"],
                "source_file": resolved_meta.get("source_file"),
                "state": row["state"],
                "custom_sections_count": row["custom_sections_count"],
                "versions": row["versions"],
            }
        )
    return rows


def get_documents_by_state(lab_id: str, states: set[str] | list[str]) -> list[dict[str, Any]]:
    states_set = set(states)
    return [row for row in get_document_list(lab_id) if row["state"] in states_set]


def get_published_documents_library() -> list[dict[str, Any]]:
    library = []
    for lab_id in LABS:
        profile, _answers, _documents = ensure_lab_initialized(lab_id)
        published_rows = [
            row for row in get_document_list(lab_id)
            if row["state"] == DOC_STATE_PUBLISHED and row["versions"]
        ]
        if not published_rows:
            continue
        library.append(
            {
                "lab_id": lab_id,
                "profile": profile,
                "documents": published_rows,
            }
        )
    return library


def get_lab_workbench_payload(lab_id: str) -> dict[str, Any]:
    documents = get_document_list(lab_id)
    return {
        "assigned_drafts": [row for row in documents if row["state"] in LAB_DRAFT_STATES],
        "submitted_documents": [row for row in documents if row["state"] == DOC_STATE_SUBMITTED],
        "published_documents": [row for row in documents if row["state"] == DOC_STATE_PUBLISHED and row["versions"]],
        "cross_lab_published": [
            row for row in get_published_documents_library()
            if row["lab_id"] != lab_id
        ],
        "all_documents": documents,
    }


def get_document_payload(lab_id: str, doc_id: str) -> dict[str, Any]:
    profile, answers, documents = ensure_lab_initialized(lab_id)
    if doc_id not in documents["documents"]:
        raise KeyError(doc_id)
    meta = get_resolved_document_meta(doc_id)
    row = documents["documents"][doc_id]
    context = _lab_context(lab_id, answers, profile)
    published_version = row["versions"][-1]["version"] if row["state"] == DOC_STATE_PUBLISHED and row["versions"] else None
    return {
        "lab": profile,
        "doc": row,
        "meta": meta,
        "context": context,
        "document_docx_view": build_generated_document_view(
            lab_id,
            doc_id,
            published_version=published_version,
        ),
        "reference_docx_view": build_idwe_reference_view(doc_id),
        "version_history": row["versions"],
    }


def update_document_draft(
    lab_id: str,
    doc_id: str,
    payload: dict[str, Any],
    *,
    actor: str,
) -> dict[str, Any]:
    documents = load_lab_documents(lab_id)
    row = documents["documents"][doc_id]
    if row["state"] not in LAB_DRAFT_STATES and row["state"] != DOC_STATE_SUBMITTED:
        raise ValueError("This document is not currently open for lab drafting.")
    sections_by_id = {
        section["section_id"]: section
        for section in row["current_draft"]["sections"]
    }
    updates = payload.get("sections", [])
    for update in updates:
        section = sections_by_id.get(update["section_id"])
        if not section:
            continue
        body = (update.get("body") or "").strip()
        section["body"] = body or section["auto_body"]
        section["manual_override"] = section["body"] != section["auto_body"]
        section["choice"] = "custom" if section["manual_override"] else section["choice"]
        section["reviewed"] = True
    row["state"] = DOC_STATE_IN_DRAFTING
    row["reviewed_at"] = _utc_now_iso()
    row["custom_sections_count"] = sum(1 for section in sections_by_id.values() if section["manual_override"])
    row["current_draft"]["updated_at"] = _utc_now_iso()
    row["current_draft"]["updated_by"] = actor
    return save_lab_documents(lab_id, documents)


def submit_document(lab_id: str, doc_id: str, *, actor_user: User) -> dict[str, Any]:
    documents = load_lab_documents(lab_id)
    row = documents["documents"][doc_id]
    if row["state"] not in LAB_DRAFT_STATES:
        raise ValueError("Only pushed or in-drafting documents can be submitted.")
    row["state"] = DOC_STATE_SUBMITTED
    row["submitted_at"] = _utc_now_iso()
    row["submitted_by"] = actor_user.username
    saved = save_lab_documents(lab_id, documents)
    notify_admins_of_submission(lab_id, doc_id, actor_user)
    return saved


def initiate_new_draft(lab_id: str, doc_id: str, *, actor_user: User) -> dict[str, Any]:
    profile, answers, documents = ensure_lab_initialized(lab_id)
    row = documents["documents"][doc_id]
    row["state"] = DOC_STATE_DRAFT_PUSHED
    row["draft_pushed_at"] = _utc_now_iso()
    row["draft_pushed_by"] = actor_user.username
    row["return_comments"] = ""
    row["return_comments_by_section"] = {}
    row["submitted_at"] = None
    row["returned_at"] = None
    row["reviewed_at"] = None
    row["current_draft"]["sections"] = build_document_sections(doc_id, _lab_context(lab_id, answers, profile))
    saved = save_lab_documents(lab_id, documents)
    notify_lab_users_of_new_draft(lab_id, doc_id, actor_user)
    return saved


def bulk_push_drafts_to_lab(lab_id: str, *, actor_user: User) -> dict[str, Any]:
    profile, answers, documents = ensure_lab_initialized(lab_id)
    pushed_doc_ids: list[str] = []
    context = _lab_context(lab_id, answers, profile)
    for meta in DOCUMENT_CATALOG:
        if meta["layer"] not in {"layer1", "placeholder"}:
            continue
        row = documents["documents"][meta["doc_id"]]
        if row["state"] in {DOC_STATE_DRAFT_PUSHED, DOC_STATE_IN_DRAFTING, DOC_STATE_SUBMITTED}:
            continue
        row["state"] = DOC_STATE_DRAFT_PUSHED
        row["draft_pushed_at"] = _utc_now_iso()
        row["draft_pushed_by"] = actor_user.username
        row["return_comments"] = ""
        row["return_comments_by_section"] = {}
        row["submitted_at"] = None
        row["returned_at"] = None
        row["reviewed_at"] = None
        row["current_draft"]["sections"] = build_document_sections(meta["doc_id"], context)
        pushed_doc_ids.append(meta["doc_id"])
    saved = save_lab_documents(lab_id, documents)
    if pushed_doc_ids:
        notify_lab_users_of_bulk_push(lab_id, pushed_doc_ids, actor_user)
    return {"documents": saved, "pushed_doc_ids": pushed_doc_ids}


def return_document_to_lab(
    lab_id: str,
    doc_id: str,
    *,
    actor_user: User,
    comments: str,
    section_comments: dict[str, str] | None = None,
) -> dict[str, Any]:
    documents = load_lab_documents(lab_id)
    row = documents["documents"][doc_id]
    row["state"] = DOC_STATE_RETURNED
    row["return_comments"] = comments.strip()
    row["return_comments_by_section"] = section_comments or {}
    row["returned_at"] = _utc_now_iso()
    row["returned_by"] = actor_user.username
    saved = save_lab_documents(lab_id, documents)
    notify_lab_users_of_return(lab_id, doc_id, actor_user)
    return saved


def publish_document(
    lab_id: str,
    doc_id: str,
    *,
    actor_user: User,
    version_mode: str,
    version_notes: str,
) -> dict[str, Any]:
    if not version_notes.strip():
        raise ValueError("Version notes are required.")
    documents = load_lab_documents(lab_id)
    row = documents["documents"][doc_id]
    new_version = increment_version(latest_version_string(row["versions"]), version_mode)
    row["versions"].append(
        {
            "version": new_version,
            "notes": version_notes.strip(),
            "published_by": actor_user.username,
            "published_at": _utc_now_iso(),
            "sections": deepcopy(row["current_draft"]["sections"]),
        }
    )
    row["state"] = DOC_STATE_PUBLISHED
    row["return_comments"] = ""
    row["return_comments_by_section"] = {}
    saved = save_lab_documents(lab_id, documents)
    notify_lab_users_of_publish(lab_id, doc_id, new_version, actor_user)
    return saved


def latest_version_string(versions: list[dict[str, Any]]) -> str:
    if not versions:
        return "0.0"
    return max((row.get("version", "0.0") for row in versions), key=_version_key)


def _version_key(version: str) -> tuple[int, int]:
    parts = str(version).split(".")
    major = int(parts[0] or 0)
    minor = int(parts[1] or 0) if len(parts) > 1 else 0
    return major, minor


def increment_version(current: str, mode: str) -> str:
    major, minor = _version_key(current)
    mode = (mode or "decimal").strip().lower()
    if mode == "major":
        return f"{major + 1}.0"
    if major == 0 and minor == 0:
        return "0.1"
    return f"{major}.{minor + 1}"


def get_lab_summary(lab_id: str) -> dict[str, Any]:
    profile, answers, documents = ensure_lab_initialized(lab_id)
    progress = progress_for_answers(answers)
    doc_rows = list(documents["documents"].values())
    draft_pushed = sum(1 for row in doc_rows if row["state"] == DOC_STATE_DRAFT_PUSHED)
    in_drafting = sum(1 for row in doc_rows if row["state"] == DOC_STATE_IN_DRAFTING)
    submitted = sum(1 for row in doc_rows if row["state"] == DOC_STATE_SUBMITTED)
    pending = sum(1 for row in doc_rows if row["state"] in {DOC_STATE_SUBMITTED, DOC_STATE_RETURNED})
    published_count = sum(1 for row in doc_rows if row["state"] == DOC_STATE_PUBLISHED and row["versions"])
    published_versions = [latest_version_string(row["versions"]) for row in doc_rows if row["versions"]]
    current_version = max(published_versions, key=_version_key) if published_versions else "0.0"
    return {
        "lab_id": lab_id,
        "profile": profile,
        "wizard_progress": progress,
        "drafts_pushed": draft_pushed,
        "drafts_in_progress": in_drafting,
        "documents_submitted": submitted,
        "documents_pending_review": pending,
        "published_documents": published_count,
        "current_version": current_version,
        "wizard_complete": answers["wizard_complete"],
    }


def list_lab_summaries() -> list[dict[str, Any]]:
    return [get_lab_summary(lab_id) for lab_id in LABS]


def infer_user_lab_ids(user: User) -> list[str]:
    candidates = {
        str(getattr(user, "username", "") or "").lower(),
        str(getattr(user, "full_name", "") or "").lower(),
        str(getattr(user, "designation", "") or "").lower(),
        str(getattr(getattr(user, "office", None), "office_name", "") or "").lower(),
        str(getattr(getattr(user, "office", None), "office_code", "") or "").lower(),
        str(getattr(getattr(user, "office", None), "location", "") or "").lower(),
    }
    joined = " ".join(text for text in candidates if text)
    matches = []
    for lab_id, definition in LABS.items():
        if any(hint in joined for hint in definition.get("hints", [])):
            matches.append(lab_id)
    return matches


def can_administer_nabl(user: User) -> bool:
    return bool(user.is_super_user() or user.is_module_admin("nabl"))


def resolve_lab_for_user(user: User, requested_lab_id: str | None = None) -> tuple[str | None, list[str]]:
    if requested_lab_id and requested_lab_id not in LABS:
        requested_lab_id = None
    if can_administer_nabl(user):
        return requested_lab_id, list(LABS.keys())
    matches = infer_user_lab_ids(user)
    if requested_lab_id and requested_lab_id in matches:
        return requested_lab_id, matches
    if requested_lab_id and not matches:
        return requested_lab_id, [requested_lab_id]
    if len(matches) == 1:
        return matches[0], matches
    return None, matches


def user_can_view_published_lab_document(user: User, lab_id: str, doc_id: str) -> bool:
    if can_administer_nabl(user):
        return True
    if lab_id not in LABS:
        return False
    documents = load_lab_documents(lab_id)
    row = documents["documents"].get(doc_id)
    if not row:
        return False
    return row["state"] == DOC_STATE_PUBLISHED and bool(row["versions"])


def get_nabl_notifications(user_id: int, limit: int = 20) -> list[Notification]:
    return (
        Notification.query
        .filter(
            Notification.user_id == user_id,
            or_(
                Notification.link.like("/nabl/%"),
                Notification.title.like("NABL %"),
            ),
        )
        .order_by(Notification.created_at.desc())
        .limit(limit)
        .all()
    )


def notify_admins_of_submission(lab_id: str, doc_id: str, actor_user: User) -> None:
    lab_name = LABS[lab_id]["display_name"]
    recipients = _nabl_admin_recipients()
    for user in recipients:
        create_notification(
            user.id,
            "NABL document submitted",
            f"{lab_name} submitted {doc_id} for review",
            severity="info",
            link=url_for("nabl.admin_document_review", lab_id=lab_id, doc_id=doc_id),
        )
    db.session.commit()


def notify_lab_users_of_return(lab_id: str, doc_id: str, actor_user: User) -> None:
    for user in _lab_user_recipients(lab_id):
        create_notification(
            user.id,
            "NABL document returned",
            f"{doc_id} returned with comments - action required",
            severity="warning",
            link=url_for("nabl.document_review", lab_id=lab_id, doc_id=doc_id),
        )
    db.session.commit()


def notify_lab_users_of_publish(lab_id: str, doc_id: str, version: str, actor_user: User) -> None:
    for user in _lab_user_recipients(lab_id):
        create_notification(
            user.id,
            "NABL document published",
            f"{doc_id} published as v{version}",
            severity="success",
            link=url_for("nabl.document_review", lab_id=lab_id, doc_id=doc_id),
        )
    db.session.commit()


def notify_lab_users_of_new_draft(lab_id: str, doc_id: str, actor_user: User) -> None:
    for user in _lab_user_recipients(lab_id):
        create_notification(
            user.id,
            "NABL draft initiated",
            f"Draft pushed for {doc_id} - ready for lab drafting",
            severity="info",
            link=url_for("nabl.document_review", lab_id=lab_id, doc_id=doc_id),
        )
    db.session.commit()


def notify_lab_users_of_bulk_push(lab_id: str, doc_ids: list[str], actor_user: User) -> None:
    label = LABS[lab_id]["display_name"]
    for user in _lab_user_recipients(lab_id):
        create_notification(
            user.id,
            "NABL draft pack pushed",
            f"{len(doc_ids)} drafts pushed to {label} for lab drafting",
            severity="info",
            link=url_for("nabl.workbench", lab_id=lab_id),
        )
    db.session.commit()


def _nabl_admin_recipients() -> list[User]:
    query = (
        User.query
        .outerjoin(ModuleAdminAssignment, ModuleAdminAssignment.user_id == User.id)
        .filter(
            User.is_active.is_(True),
            or_(
                User.role.has(name="superuser"),
                ModuleAdminAssignment.module_code == "nabl",
            ),
        )
        .distinct()
    )
    return query.all()


def _lab_user_recipients(lab_id: str) -> list[User]:
    recipients = []
    for user in User.query.filter(User.is_active.is_(True)).all():
        resolved_lab, _matches = resolve_lab_for_user(user)
        if resolved_lab == lab_id and not can_administer_nabl(user):
            recipients.append(user)
    return recipients


def build_docx_bytes(
    lab_id: str,
    doc_id: str,
    *,
    published_version: str | None = None,
) -> bytes:
    if Document is None:  # pragma: no cover
        raise RuntimeError("python-docx is required to generate NABL documents.")
    profile, answers, documents = ensure_lab_initialized(lab_id)
    payload = documents["documents"][doc_id]
    meta = get_resolved_document_meta(doc_id)
    sections = payload["current_draft"]["sections"]
    if published_version:
        for version in payload["versions"]:
            if version["version"] == published_version:
                sections = version["sections"]
                break
    document = Document()
    _apply_doc_styles(document)
    _add_logo(document)
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(f"{doc_id}\n{meta['title']}")
    run.bold = True
    run.font.size = Pt(16)
    subheader = document.add_paragraph()
    subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subheader.add_run(
        f"{profile['display_name']} | {profile['location']} | Approved by {profile['approved_by']}"
    ).italic = True
    if meta.get("source_file"):
        source_note = document.add_paragraph()
        source_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        source_note.add_run(f"Aligned to IDWE master source: {meta['source_file']}").italic = True
    document.add_paragraph("")
    for section in sections:
        document.add_heading(section["title"], level=2)
        source = section.get("source_question")
        if source:
            document.add_paragraph(f"Source question: {source}")
        for block in section["body"].split("\n"):
            document.add_paragraph(block)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_cover_docx_bytes(lab_id: str) -> bytes:
    if Document is None:  # pragma: no cover
        raise RuntimeError("python-docx is required to generate NABL documents.")
    summary = get_lab_summary(lab_id)
    profile = summary["profile"]
    version = summary["current_version"]
    version_label = (
        f"Version {version} - For NABL Submission"
        if _version_key(version) > (0, 0)
        else "Working Draft v0.2"
    )
    document = Document()
    _apply_doc_styles(document)
    _add_logo(document)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Application for NABL Accreditation")
    title_run.bold = True
    title_run.font.size = Pt(20)
    for line in [profile["display_name"], profile["location"], datetime.now().strftime("%d %b %Y"), version_label]:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(line)
    document.add_page_break()
    document.add_heading("Table of Contents", level=1)
    compile_docs = [get_resolved_document_meta(row["doc_id"]) for row in DOCUMENT_CATALOG if row["compile"]]
    for index, meta in enumerate(compile_docs, start=1):
        document.add_paragraph(f"{index:02d}. {meta['doc_id']} - {meta['title']}")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_merged_docx_bytes(lab_id: str) -> bytes:
    if Document is None:  # pragma: no cover
        raise RuntimeError("python-docx is required to generate NABL documents.")
    profile, _answers, documents = ensure_lab_initialized(lab_id)
    document = Document()
    _apply_doc_styles(document)
    _add_logo(document)
    document.add_heading(f"{profile['display_name']} - NABL Application", level=0)
    compile_docs = [get_resolved_document_meta(row["doc_id"]) for row in DOCUMENT_CATALOG if row["compile"]]
    for index, meta in enumerate(compile_docs):
        payload = documents["documents"][meta["doc_id"]]
        document.add_page_break()
        document.add_heading(f"{meta['doc_id']} - {meta['title']}", level=1)
        for section in payload["current_draft"]["sections"]:
            document.add_heading(section["title"], level=2)
            for block in section["body"].split("\n"):
                document.add_paragraph(block)
        if index == len(compile_docs) - 1:
            break
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def compile_application_zip(lab_id: str) -> tuple[io.BytesIO, str]:
    summary = get_lab_summary(lab_id)
    version = summary["current_version"]
    filename = f"{lab_id}_NABL_Application_v{version}.zip"
    stream = io.BytesIO()
    with zipfile.ZipFile(stream, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("00_Cover_and_TOC.docx", build_cover_docx_bytes(lab_id))
        compile_docs = [get_resolved_document_meta(meta["doc_id"]) for meta in DOCUMENT_CATALOG if meta["compile"]]
        for index, meta in enumerate(compile_docs, start=1):
            archive.writestr(
                f"{index:02d}_{meta['doc_id']}.docx",
                build_docx_bytes(lab_id, meta["doc_id"]),
            )
        archive.writestr(
            f"{lab_id}_NABL_Application_MERGED_v{version}.docx",
            build_merged_docx_bytes(lab_id),
        )
    stream.seek(0)
    return stream, filename


def _apply_doc_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)


def _add_logo(document: Document) -> None:
    for path in LOGO_CANDIDATES:
        if path.exists():
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(str(path), width=Inches(1.25))
            return
