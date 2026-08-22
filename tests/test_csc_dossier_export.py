from io import BytesIO

from docx import Document
from docx.oxml.ns import qn

from app.core.services.csc_dossier_export import build_enterprise_dossier


def _context():
    return {
        "draft": {
            "spec_number": "ONGC/WS/27/2015",
            "chemical_name": "LOW STRENGTH PROPPANT",
            "material_code": "090001039 (LSP 12/20)",
            "test_procedure": "ONGC/ Test Procedure Vol-III / WS / 27",
            "version_display": "v0",
        },
        "summary_rows": [
            {"label": "Specification", "value": "ONGC/WS/27/2015"},
            {"label": "Chemical", "value": "LOW STRENGTH PROPPANT"},
            {"label": "Category", "value": "Well Stimulation Chemicals"},
            {"label": "Version", "value": "v0"},
            {"label": "Impact Classification", "value": "LOW IMPACT"},
            {"label": "Prepared By", "value": "Specification Owner"},
            {"label": "Reviewed By", "value": ""},
            {"label": "Last Revised", "value": "21 Aug 2026"},
            {"label": "On Corporate Register", "value": "Yes"},
            {"label": "Standard Testing Time", "value": "1 day"},
        ],
        "covered_chemicals": [
            {"chemical_name": "Low Strength Proppant (12/20)", "material_code": "90001039", "standard_days": 1},
            {"chemical_name": "Low Strength Proppant (20/40)", "material_code": "100101589", "standard_days": None},
            {"chemical_name": "Low Strength Proppant (40/70)", "material_code": "100000281", "standard_days": None},
        ],
        "parameter_rows": [
            {
                "parameter_name": "Physical State",
                "parameter_type": "Essential",
                "final_requirement": "The material shall be in the form of single grains",
                "unit_of_measure": "—",
                "conditions": "—",
                "procedure_text": "—",
            },
            {
                "parameter_name": "Bulk Density",
                "parameter_type": "Vital",
                "final_requirement": "1.55 maximum",
                "unit_of_measure": "g/cc",
                "conditions": "At 25 °C",
                "procedure_text": "ASTM method",
            },
        ],
        "section_rows": [
            {"label": "Background", "text": "Legacy specification", "source_text": "—"},
            {"label": "Recommendation", "text": "Retain for corporate use", "source_text": "—"},
        ],
        "master_rows": [{"label": "Material Code", "value": "090001039", "source_value": "—"}],
        "material_property_rows": [],
        "storage_rows": [{"label": "Storage", "value": "Store dry", "source_value": "—"}],
        "impact_rows": [{"label": "Impact Classification", "value": "LOW IMPACT", "source_value": "—"}],
        "issue_rows": [{"label": "Quality", "is_present": False, "note": ""}],
        "versions": [
            {
                "version": "0",
                "created_at": "10 Mar 2026",
                "created_by": "Specification Owner",
                "action": "RESET_BASELINE",
                "reason": "Current content preserved as v0",
                "is_current": True,
            }
        ],
        "latest_review_notes": "",
    }


def test_enterprise_dossier_has_controlled_sections_without_empty_legacy_comparisons():
    document = Document(BytesIO(build_enterprise_dossier(_context())))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)

    assert document.core_properties.category == "Controlled Technical Document"
    assert "CORPORATE SPECIFICATION DOSSIER" in text
    assert "2. Specification Requirements" in text
    assert "6. Governance and Revision Record" in text
    assert "Low Strength Proppant (12/20)" in table_text
    assert "Physical State" in table_text
    assert "Published Baseline" not in table_text
    assert document.settings.element.find(qn("w:evenAndOddHeaders")) is None
    assert all(section.different_first_page_header_footer for section in document.sections)
    header_text = "\n".join(
        cell.text
        for table in document.sections[0].first_page_header.tables
        for row in table.rows
        for cell in row.cells
    )
    footer_text = "\n".join(
        cell.text
        for table in document.sections[0].first_page_footer.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "CORPORATE SPECIFICATION DOSSIER" in header_text
    assert "CONTROLLED COPY" in footer_text


def test_enterprise_dossier_tables_use_fixed_full_width_geometry():
    document = Document(BytesIO(build_enterprise_dossier(_context())))
    for table in document.tables:
        table_width = table._tbl.tblPr.find(qn("w:tblW"))
        assert table_width is not None
        assert table_width.get(qn("w:w")) == "9360"
        assert table_width.get(qn("w:type")) == "dxa"
